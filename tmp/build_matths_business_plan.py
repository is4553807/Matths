from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/sangyoonlee/Desktop/SangYoon Lee/SINGAPORE 2025-/Personal Projects/Matths")
OUT = ROOT / "output" / "docs" / "Matths_사업소개서_및_공동사업_운영계획서_v1.1.docx"

# User-requested typography: Malgun Gothic across Korean and Latin runs.
FONT_LATIN = "Malgun Gothic"
FONT_KO = "Malgun Gothic"
# Monochrome formal-document override of standard_business_brief.
BLUE = "111111"
DARK_BLUE = "111111"
INK = "111111"
TEXT = "1A1A1A"
MUTED = "5F5F5F"
LIGHT = "F2F2F2"
PALE_BLUE = "E8E8E8"
CALLOUT = "F7F7F7"
BORDER = "B7B7B7"
WHITE = "FFFFFF"
RISK = "333333"
GOLD = "555555"

PAGE_DXA = 9360
TABLE_INDENT = 120


def set_font(run, size=None, bold=None, italic=None, color=TEXT, name=FONT_LATIN):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), FONT_KO)
    fonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], indent=TABLE_INDENT):
    if sum(widths) != PAGE_DXA:
        raise ValueError(f"table widths must total {PAGE_DXA}, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[i] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_paragraph_border(paragraph, color=BORDER, size=6, side="bottom"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "5")
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_KO)
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    for name, size, color, italic in (
        ("Table Source", 8.5, MUTED, False),
        ("Source List", 8.1, TEXT, False),
        ("Small Note", 9.0, MUTED, False),
        ("Kicker", 9.0, BLUE, False),
    ):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10
    styles["Source List"].paragraph_format.left_indent = Inches(0.32)
    styles["Source List"].paragraph_format.first_line_indent = Inches(-0.32)
    styles["Source List"].paragraph_format.space_before = Pt(1)
    styles["Source List"].paragraph_format.space_after = Pt(2)
    styles["Source List"].paragraph_format.line_spacing = 1.02
    styles["Kicker"].font.bold = True
    styles["Kicker"].font.all_caps = True


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    # Add explicit preset-matching definitions; paragraphs link to these IDs.
    def add_abstract(abstract_id, num_fmt, text, left=720, hanging=360):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_id_el = OxmlElement("w:abstractNumId")
        abstract_id_el.set(qn("w:val"), str(abstract_id))
        num.append(abstract_id_el)
        numbering.append(num)

    add_abstract(910, "bullet", "•")
    add_abstract(911, "decimal", "%1.")
    add_num(910, 910)
    add_num(911, 911)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_para(doc, text="", bold_prefix=None, align=None, style=None, after=None, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if keep:
        p.paragraph_format.keep_together = True
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.10
        apply_num(p, 910)
        set_font(p.add_run(item))


def add_numbered(doc, items: Iterable[str]):
    # Each list gets a fresh concrete numbering instance so it restarts at 1.
    num_id = getattr(doc, "_matths_next_num_id", 912)
    numbering = doc.part.numbering_part.element
    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), str(num_id))
    abstract_id_el = OxmlElement("w:abstractNumId")
    abstract_id_el.set(qn("w:val"), "911")
    num_el.append(abstract_id_el)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num_el.append(lvl_override)
    numbering.append(num_el)
    doc._matths_next_num_id = num_id + 1
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.10
        apply_num(p, num_id)
        set_font(p.add_run(item))


def add_callout(doc, label, text, fill=CALLOUT, border=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, border, size=12, side="left")
    set_font(p.add_run(label + "  "), size=10, bold=True, color=border)
    set_font(p.add_run(text), size=10, color=TEXT)
    return p


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int],
              aligns: Sequence[int] | None = None, header_fill=LIGHT, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(header), size=9.2, bold=True, color=INK)
    for row_vals in rows:
        row = table.add_row()
        for j, value in enumerate(row_vals):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.alignment = aligns[j] if aligns else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.08
            set_font(p.add_run(str(value)), size=font_size, color=TEXT)
    # New rows need geometry too.
    set_table_geometry(table, widths)
    for row in table.rows:
        prevent_row_split(row)
    return table


def add_table_source(doc, text):
    p = doc.add_paragraph(style="Table Source")
    set_font(p.add_run(text), size=8.5, color=MUTED)
    return p


def add_section_heading(doc, number, title, subtitle=None):
    p = doc.add_paragraph(style="Kicker")
    p.paragraph_format.keep_with_next = True
    label = "부록" if number == "A" else f"제{int(number)}장"
    set_font(p.add_run(label), size=9, bold=True, color=MUTED)
    h = doc.add_paragraph(style="Heading 1")
    set_font(h.add_run(title), size=16, bold=True, color=INK)
    if subtitle:
        p2 = doc.add_paragraph(style="Small Note")
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.keep_with_next = True
        set_font(p2.add_run(subtitle), size=9, color=MUTED)


def new_page(doc):
    # Every major section starts on a fresh page. This removes ambiguous
    # cross-page section boundaries in administrative review copies.
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    set_font(p.add_run(""), size=1, color=WHITE)


def hard_new_page(doc):
    # Use only at true form-factor boundaries such as after the cover.
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    set_font(p.add_run(""), size=1, color=WHITE)


def add_source(doc, code, title, org, date, url, note=""):
    p = doc.add_paragraph(style="Source List")
    set_font(p.add_run(f"[{code}] {org}, {title}"), size=8.1, bold=True, color=TEXT)
    set_font(p.add_run(f" ({date}). "), size=8.1, color=MUTED)
    add_hyperlink(p, url, url)
    if note:
        set_font(p.add_run(f" — {note}"), size=8.1, color=MUTED)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False

    configure_styles(doc)
    configure_numbering(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(4)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_font(hp.add_run("MATTHS  |  사업소개서"), size=8.5, bold=True, color=MUTED)
    set_font(hp.add_run("\t내부검토용 · 대외배포금지"), size=8.5, bold=True, color=INK)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(fp.add_run("Matths  |  v1.1  |  "), size=8.5, color=MUTED)
    add_page_field(fp)

    doc.core_properties.title = "Matths 사업소개서 및 공동사업 운영계획서"
    doc.core_properties.subject = "공동사업자 등록 및 내부 협의용"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "Matths, 사업소개서, 공동사업, 교육서비스, GOAT Arena"
    return doc


def build():
    doc = setup_document()

    # COVER — monochrome proposal_centerpiece using standard_business_brief.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("MATTHS"), size=12, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("공동사업자 등록·협의용"), size=9, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("사업소개서 및 공동사업 운영계획서"), size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    set_font(p.add_run("시각화 기반 고등수학 학습 플랫폼 · GOAT Arena"), size=13, color=MUTED)

    cover_rows = [
        ("문서 목적", "공동사업자 등록·협의 및 사업성 검토"),
        ("사업 형태", "개인 공동사업(예정)"),
        ("작성 기준일", "2026년 8월 2일"),
        ("문서 버전", "v1.1"),
    ]
    table = add_table(doc, ["구분", "내용"], cover_rows, [2200, 7160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], header_fill=PALE_BLUE, font_size=9.6)
    for c in table.rows[0].cells:
        set_cell_shading(c, PALE_BLUE)
    add_table_source(doc, "주: 상호, 사업장 소재지, 공동사업자 인적사항 및 손익분배비율은 사업자등록 신청서·동업계약서에서 별도 확정한다.")

    add_callout(
        doc,
        "문서 등급",
        "본 문서는 공동사업자 제출 및 내부 의사결정을 위한 비공개 자료이다. 이용자 홍보물, 투자권유 자료 또는 세무·법률 의견서로 사용하지 않는다.",
        fill=CALLOUT,
        border=INK,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    set_font(p.add_run("공동사업자 등록 및 내부 사업성 검토를 위한 공식 문서"), size=9, color=MUTED)

    hard_new_page(doc)
    add_section_heading(doc, "00", "요약", "공동사업자 검토를 위한 핵심 사업성·수익성·세무 판단 요약")
    add_callout(
        doc,
        "핵심 판단",
        "Matths는 2026년 운영 기준 개편 수능에 대응하는 2022 개정 고등학교 수학 교육과정 13개 과목을 제공하는 완성형 플랫폼이다. 실제 이용자는 학생이지만 결제와 선택은 학부모·보호자·제3자가 담당할 수 있으므로, 학생의 학습경험과 구매자의 신뢰·성과확인을 동시에 설계하는 비대칭 시장 구조를 전제로 한다.",
    )

    snapshot_rows = [
        ("사업 영역", "온라인 고등수학 교육·학습평가 플랫폼"),
        ("이용자·구매자", "실제 이용자: 학생 / 결제·의사결정자: 학부모·보호자·제3자"),
        ("교육과정 범위", "공통·일반선택·진로선택·융합선택을 포함한 고등학교 수학 13개 과목"),
        ("핵심 제품", "시각화 학습, 문제·오답 복습, 평가, 진도·성과관리, GOAT Arena"),
        ("기준 가격", "29일 학습 패키지 29,000원"),
        ("수익 구조", "학부모·보호자·제3자의 학습권 결제 + 학생 이용 + 조건부 페이백·재구독"),
        ("기준 손익", "최초 결제자 1,000명 기준 연 결제 1,940건, 고객 결제총액 약 5,626만 원"),
        ("문서상 결론", "구매자에게는 가격·성과 투명성, 학생에게는 학습효과·동기·보상 경험을 제공하는 다중 고객 모델"),
    ]
    add_table(doc, ["항목", "요약"], snapshot_rows, [2200, 7160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.0)

    doc.add_paragraph(style="Heading 2").add_run("의사결정 전제")
    add_bullets(doc, [
        "현재 손익은 실제 유료 운영 실적이 아닌 Monte Carlo 기반 정책 검증 추정치다.",
        "연매출은 세액감면으로 증가하지 않는다. 세액감면은 세후 유보금·현금흐름을 증가시킨다.",
        "29,000원이 부가가치세 포함 소비자가격인지, 교육용역 면세가 가능한지가 회계상 매출과 이익을 크게 바꾼다.",
        "청년창업 세액감면은 청년 여부만으로 적용되지 않으며, 실제 주업종·창업일·사업장 위치·공동사업 손익분배비율을 함께 충족해야 한다.",
    ])

    new_page(doc)
    add_section_heading(doc, "01", "사업 개요", "공식 사업계획서의 문제인식·실현가능성·성장전략·팀 구성 체계를 반영 [S4]")
    doc.add_paragraph(style="Heading 2").add_run("1.1 사업 정의")
    add_para(doc, "Matths는 수학 공식을 정답으로 제시하는 데 그치지 않고, 도형·그래프·식이 변하는 과정을 장면 단위로 보여주는 고등수학 학습 플랫폼이다. 2026년 운영 기준으로 2022 개정 교육과정의 공통수학1·2부터 일반선택·진로선택·융합선택까지 총 13개 과목을 제공하며, 개편 수능과 학교별 과목 편성에 대응한다. 학생은 개념의 생성 원리를 시각적으로 이해한 뒤 문제에 적용하고, 오답이 발생하면 막힌 단계로 돌아가 재학습한다. 서비스의 주된 급부는 교육 콘텐츠·문제 풀이·학습 기록·평가·피드백이며, GOAT Arena는 학습 성과를 경쟁과 랭킹으로 확장하는 부가 기능이다. [I1][I2][S12]")

    doc.add_paragraph(style="Heading 2").add_run("1.2 해결하려는 문제")
    problem_rows = [
        ("원리 단절", "공식과 정답을 먼저 제시해 개념이 만들어지는 과정이 생략됨", "도형·그래프·식의 단계별 변화를 함께 제시"),
        ("오답 반복", "어디서 막혔는지 모른 채 전체 강의나 해설을 반복함", "개념·조건 해석·적용·계산 단계별 복귀"),
        ("학습 지속성", "단기 집중 후 이탈하고 성취 확인이 늦음", "일일 학습·주간 평가·랭킹·페이백으로 동기 강화"),
        ("가격 부담", "참여 고등학생의 사교육비가 월평균 79.3만 원 수준", "29,000원 진입 가격으로 보조 학습재 포지셔닝"),
    ]
    add_table(doc, ["문제", "현재 불편", "Matths 접근"], problem_rows, [1500, 3780, 4080], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.9)

    doc.add_paragraph(style="Heading 2").add_run("1.3 사업 목표")
    add_numbered(doc, [
        "2022 개정 고등학교 수학 교육과정 13개 과목 전체에 시각화 학습·문제 적용·오답 복습·평가의 완결된 흐름을 제공한다.",
        "유료 학습 패키지와 주간 평가를 결합해 반복 이용형 학습 운영 모델을 확립하고, 상용 운영 데이터로 수익 가정을 정기 보정한다.",
        "페이백과 Division 경쟁을 교육 성취 보상으로 운영하되, 결제·학습권·랭킹·환급 원장을 분리하여 공정성과 재현성을 확보한다.",
    ])

    new_page(doc)
    add_section_heading(doc, "02", "시장 배경 및 비대칭 고객 구조", "결제자와 실제 이용자가 분리되는 교육시장의 특성을 반영한 다중 고객 전략")
    doc.add_paragraph(style="Heading 2").add_run("2.1 2025년 사교육 시장")
    add_para(doc, "교육부·국가데이터처의 2025년 초중고 사교육비 조사에 따르면, 고등학교 사교육비 총액은 약 7.8조 원, 고등학생 사교육 참여율은 63.0%이다. 사교육 참여 고등학생 1인당 월평균 지출은 79만 3천 원이며, 고1은 80만 6천 원으로 가장 높다. [S1][S2]")
    market_rows = [
        ("고등학교 사교육비 총액", "약 7.8조 원", "전년 대비 4.3% 감소"),
        ("고등학생 사교육 참여율", "63.0%", "전체 고등학생 중 참여 비중"),
        ("참여 고등학생 월평균 지출", "793,000원", "2025년 학교급 평균"),
        ("고1 참여학생 월평균 지출", "806,000원", "학년별 최고 수준"),
        ("Matths 기준 가격", "29,000원", "참여학생 월평균 지출의 약 3.7%"),
    ]
    add_table(doc, ["지표", "2025년 값", "해석"], market_rows, [2900, 1800, 4660], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.2)
    add_table_source(doc, "자료: 국가데이터처·교육부, 2025년 초중고사교육비조사 결과(2026.3.12 공표) [S1][S2]. 금액 반올림.")
    add_callout(doc, "시장 해석", "Matths의 29,000원은 기존 학원·과외를 전면 대체하기보다 월 사교육 예산의 약 3.7%로 추가 채택 가능한 보조 학습 서비스 가격대다. 따라서 초기 전략은 ‘대체재’보다 ‘개념 이해와 학습 지속성을 보완하는 디지털 학습재’가 적합하다.")

    hard_new_page(doc)
    doc.add_paragraph(style="Heading 2").add_run("2.2 비대칭 고객 구조")
    add_para(doc, "Matths의 시장 비대칭성은 서비스 이용과 비용 부담의 주체가 일치하지 않을 수 있다는 점에 있다. 학생은 학습경험과 성취를 직접 소비하지만, 학부모·보호자 또는 제3자가 가격을 비교하고 결제하며 이용 지속 여부에 영향을 미친다. 따라서 학생만을 최종 고객으로 보거나 결제자만을 구매 고객으로 보는 단일 고객 접근은 적절하지 않다.")
    customer_rows = [
        ("실제 이용자", "학생·N수생", "이해도·성취·동기·편의", "시각화 학습, 오답 복귀, 진도, 평가, Arena"),
        ("구매·결제자", "학부모·보호자", "가격·신뢰·성과·안전", "성과 요약, 이용기간·결제·환급조건의 투명성"),
        ("제3자 구매자", "친족·후원자·장학주체·교육기관", "학습권 제공·대상 확인", "대상 학생 지정, 사용·성과 확인, 증빙 가능한 결제"),
        ("영향자", "교사·상담자·친구", "추천·학습방향·사회적 신뢰", "교육과정 정보, 공유 가능한 성과 요약, 검증된 운영"),
    ]
    add_table(doc, ["역할", "대상", "핵심 판단기준", "제공 가치"], customer_rows, [1500, 2200, 2300, 3360], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.7)

    doc.add_paragraph(style="Heading 2").add_run("2.3 다중 고객 접근전략")
    strategy_rows = [
        ("학생 획득", "시각화 학습·무료 체험·랭킹 경험", "가입·체험·활성학습·완료율"),
        ("학부모 전환", "가격 대비 가치·성과 리포트·안전한 결제", "결제전환·갱신·환불·문의"),
        ("제3자 구매", "학습권 선물·후원·대상 학생 지정", "구매자-학생 연결·사용개시·재구매"),
        ("유지·확장", "학생 이용성과와 구매자 만족을 분리 관리", "학생 리텐션·구매자 리텐션·추천"),
    ]
    add_table(doc, ["목적", "핵심 수단", "관리지표"], strategy_rows, [1900, 4400, 3060], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.9)

    doc.add_paragraph(style="Heading 2").add_run("2.4 경쟁 구도")
    add_bullets(doc, [
        "학원·과외: 높은 상호작용과 관리 강점이 있으나 비용과 장소 제약이 큼.",
        "인터넷 강의: 범위와 강사 선택이 넓으나 수동 시청과 오답 구간 복귀의 한계가 있음.",
        "문제집·문제은행: 반복 연습에 강하나 공식이 생성되는 원리를 동적으로 설명하기 어려움.",
        "AI 풀이 서비스: 즉시 해설은 빠르지만 교육과정 기반의 누적 학습·공정한 경쟁·환급 운영은 별도 설계가 필요함.",
    ])

    new_page(doc)
    add_section_heading(doc, "03", "서비스 구성 및 차별성", "2022 개정 고등학교 수학 전 과목과 경쟁 시스템을 하나의 완성형 서비스로 연결")
    doc.add_paragraph(style="Heading 2").add_run("3.1 교육과정 전체 범위")
    add_para(doc, "Matths는 공통수학 일부가 아니라 2026년부터 순차 적용되는 개편 수능 체제에 대응하여 2022 개정 고등학교 수학 교육과정의 13개 과목 전체를 서비스 범위로 한다. 과목별 개념 시각화·문제·평가·오답 복습이 동일한 학습 체계 안에서 제공된다. [S12][I1]")
    curriculum_rows = [
        ("공통과목", "2", "공통수학1, 공통수학2"),
        ("일반선택", "3", "대수, 미적분Ⅰ, 확률과 통계"),
        ("진로선택", "5", "기하, 미적분Ⅱ, 경제 수학, 인공지능 수학, 직무 수학"),
        ("융합선택", "3", "실용 통계, 수학과 문화, 수학과제 탐구"),
    ]
    add_table(doc, ["교육과정 구분", "과목 수", "제공 과목"], curriculum_rows, [1900, 1300, 6160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.0)

    doc.add_paragraph(style="Heading 2").add_run("3.2 핵심 학습 흐름")
    flow_rows = [
        ("1", "개념 시각화", "공식이 만들어지는 과정을 도형·그래프·식의 변화로 제시"),
        ("2", "문제 적용", "학습한 개념을 교육과정·모의고사 유형 문제에 적용"),
        ("3", "막힌 지점 진단", "개념·조건 해석·적용·계산 중 실패 단계를 구분"),
        ("4", "모션그래픽 오답", "틀린 문제의 조건과 풀이를 장면 단위로 재설명"),
        ("5", "유사 문제 재도전", "새 문제로 실제 이해 여부를 재검증"),
        ("6", "주간 평가·랭킹", "공식 모의고사와 Arena 경쟁으로 성취를 지속 측정"),
    ]
    add_table(doc, ["단계", "기능", "이용자 가치"], flow_rows, [850, 2200, 6310], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.1)
    add_table_source(doc, "내부 서비스 소개 및 구현 문서 기준 [I1][I2].")

    doc.add_paragraph(style="Heading 2").add_run("3.3 제품 차별성")
    diff_rows = [
        ("시각 설명", "한 화면에 하나의 핵심 변화를 배치해 공식·도형·그래프를 연결"),
        ("오답 복귀", "전체 강의 반복이 아니라 사용자가 막힌 장면부터 재학습"),
        ("평가 분리", "시험 기반 실력지표와 Arena 경쟁지표를 분리해 결과 왜곡을 방지"),
        ("정책 버전", "가격·페이백·랭킹 규칙을 주기 시작 시점에 고정하고 소급 변경하지 않음"),
        ("감사 가능성", "결제·학습권·경기·정산 이벤트를 원장으로 남겨 결과를 재현"),
        ("리텐션 설계", "학습권 만료·주간 평가·Division 이동·재구독을 하나의 운영 주기로 연결"),
    ]
    add_table(doc, ["차별 요소", "구현 방향"], diff_rows, [2200, 7160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.2)

    doc.add_paragraph(style="Heading 2").add_run("3.4 서비스 구축 상태")
    add_para(doc, "본 문서는 제품 개발이 완료되고 상용 운영이 가능한 상태를 전제로 작성한다. 13개 과목의 콘텐츠·문제·평가, 회원·진도·오답·성과관리, Sub/Main Division, 결제·환불·페이백 원장, 관리자·보안·개인정보·CS 및 상용 인프라가 모두 구축·검증된 것으로 본다. 다만 이후 재무수치는 제품 완성도와 별개로 실제 매출 실적을 반영하지 않은 예측치다. [I1][I2]")
    status_rows = [
        ("교육과정·콘텐츠", "완료", "13개 과목의 개념·문제·평가·오답 학습 제공"),
        ("학습·평가", "완료", "시각화·진도·성과 리포트·주간 평가 운영"),
        ("결제·정산", "완료", "PG, 환불, 페이백 심사·지급·원장 관리"),
        ("Arena·랭킹", "완료", "Sub/Main 규칙, 서버 판정, 감사기록, 이상 징후 통제"),
        ("운영·관리", "완료", "관리자, 모니터링, 보안·개인정보, CS, 백업 체계"),
        ("매출·수익 데이터", "예측", "완성 제품 기준 예측이며 실제 상용 실적은 미반영"),
    ]
    add_table(doc, ["영역", "상태", "판단"], status_rows, [1900, 1600, 5860], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.0)

    new_page(doc)
    add_section_heading(doc, "04", "비대칭 시장 구조와 수익모델", "학습권의 구매자와 실제 이용자를 분리하고 양측의 가치를 함께 관리하는 구조")
    doc.add_paragraph(style="Heading 2").add_run("4.1 비대칭 시장 구조")
    add_para(doc, "Matths에서는 학부모·보호자·친족·후원자·교육기관 등 구매자가 학습권을 결제하고, 지정된 학생이 서비스를 실제 이용할 수 있다. 따라서 계약·결제·환불 안내는 구매자에게, 학습·평가·랭킹 경험은 학생에게 제공하되, 개인정보와 동의 범위 안에서 구매자에게 이용·성과 요약을 제공하는 이중 가치체계를 적용한다.")
    asym_rows = [
        ("구매자", "가격·안전·성과 확인·결제 통제", "투명한 가격·환급조건, 성과 요약, 갱신·환불 관리"),
        ("학생", "이해·성취·동기·사용 편의", "전 과목 시각화 학습, 오답 복귀, 평가, Arena"),
        ("플랫폼", "권리·데이터·책임의 분리", "구매자-학생 연결, 학습권 배정, 동의·접근권한 관리"),
    ]
    add_table(doc, ["주체", "핵심 가치", "플랫폼 대응"], asym_rows, [1800, 3100, 4460], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.0)

    doc.add_paragraph(style="Heading 2").add_run("4.2 수익모델 및 결제 흐름")
    add_para(doc, "구매자는 지정 학생에게 적용되는 29일 학습 패키지를 29,000원에 결제한다. 학생이 Sub Division에서 사전 고지된 연속 학습·경기·점수·무결성 조건을 충족하면 결제대금의 일부 또는 전부가 결제·환급 정책에 따라 처리된다. Main Division은 추가 페이백 없이 추가 학습일수와 상위 랭킹 경쟁을 제공해 학생 체류와 구매자의 재구매 동기를 강화한다. [I3][I4]")
    bm_rows = [
        ("구매", "학생 지정형 학습권", "학부모·보호자·제3자가 결제"),
        ("배정·이용", "29일 학습·평가", "지정 학생에게 학습권을 배정하고 전 과목 서비스 제공"),
        ("성과", "Sub 페이백", "조건 충족자에 한해 50%·80%·100% 환급"),
        ("유지", "Main 경쟁", "페이백 종료 후 추가 학습일수·랭킹 경쟁"),
        ("재구매", "만료 후 재구독", "72시간 내 복귀 혜택과 이후 재배치"),
    ]
    add_table(doc, ["단계", "수익·운영 장치", "사업 효과"], bm_rows, [1200, 2700, 5460], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.1)

    doc.add_paragraph(style="Heading 2").add_run("4.3 단위경제성과 조건부 페이백")
    add_para(doc, "현재 내부 시뮬레이션에서 페이백 수령자는 결제자의 약 9.81%, 전체 매출 대비 지급액은 약 4.95%로 추정된다. 모든 결제자는 교육 서비스를 제공받지만 현금 환급은 별도의 높은 성취 조건을 충족한 일부에게만 발생한다. 이에 따라 사용자에게는 ‘학습 서비스 + 성취형 환급 기회’가 제공되고, 운영자에게는 환급 재원을 부담한 뒤에도 높은 평균 공헌이익이 남는다. [I4]")
    unit_rows = [
        ("결제 1건", "29,000원", "고객 결제총액"),
        ("예상 페이백", "약 1,436원", "매출의 4.95%"),
        ("PG 변동비", "약 1,085원", "가정 수수료 3.74%"),
        ("PG 고정비 배분", "약 330원", "1,000건 기준 내부 가정"),
        ("직접 공헌이익", "약 26,150원", "변동 운영비 전, 결제액의 약 90.2%"),
        ("운영비 5,000원 반영", "약 21,150원", "결제액의 약 72.9%"),
    ]
    add_table(doc, ["항목", "1건당 금액", "해석"], unit_rows, [2600, 2000, 4760], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.2)
    add_table_source(doc, "내부 손익 시뮬레이션의 30일 결제 1,000건 평균값을 1건 기준으로 환산 [I4].")

    add_callout(doc, "운영 원칙", "페이백은 비대칭 시장 구조의 정의가 아니라 학습성과 보상 장치다. 구매자에게 수령 확률·조건·검증·지급 절차를 결제 전에 명확히 고지하고, 학생에게는 학습조건과 성취기준을 이해 가능한 방식으로 안내한다. 환급 난도를 은폐하거나 이용자 손실을 전제로 수익을 설명하지 않는다.", fill=CALLOUT, border=INK)

    doc.add_paragraph(style="Heading 2").add_run("4.4 손익에 가장 민감한 변수")
    add_bullets(doc, [
        "Sub 만료자와 Main 강등자의 24시간·72시간 내 재구독률",
        "실제 페이백 수령률 및 50%·80%·100% 구간별 지급 분포",
        "문제 제작·검수·고객지원·서버의 결제 1건당 변동원가",
        "신규 사용자 획득비용(CAC)과 유료 전환율",
        "Main 경기의 학습일수 이전이 평균 체류기간에 미치는 영향",
        "부가가치세 과·면세 및 청년창업 세액감면 적용 여부",
    ])

    new_page(doc)
    add_section_heading(doc, "05", "Main Division 운영 규칙(요약)", "사용자용 상세 룰이 아닌 공동사업자·행정 제출용 핵심 운영 원칙")
    add_callout(doc, "제도 목적", "Main Division은 Sub Division에서 페이백 및 진입 자격을 달성한 이용자가 추가 학습일수와 상위 Arena 지위를 두고 경쟁하는 상위 단계다. Main에서는 새로운 페이백을 지급하지 않는다. [I3]")

    doc.add_paragraph(style="Heading 2").add_run("5.1 핵심 규칙")
    main_rows = [
        ("진입·초기값", "Sub 페이백·진입 자격 달성 후 이동하며, 시작 학습일수는 확정 산식으로 부여"),
        ("학습일수", "사용 가능·초대 예약·경기 잠금으로 구분하고 매일 1일 차감"),
        ("매칭", "이용자는 목표 티어만 선택하며 상대는 서버가 적격 후보 중 무작위 배정"),
        ("경기·정산", "티어 차이에 따른 허용 범위와 최소 잔액을 적용하고 서버 판정으로 일괄 정산"),
        ("페이백", "Main에서는 추가 페이백을 지급하지 않고 학습일수와 랭킹 경쟁만 제공"),
        ("만료·재구독", "학습일수 소진 시 Sub로 이동하며, 정해진 기한 내 재결제에 복귀 기준 적용"),
        ("주간 동결", "마감 이후 신규 경기를 차단하고 결과를 동결한 뒤 정해진 시각에 공개"),
    ]
    add_table(doc, ["구분", "공식 요약"], main_rows, [2200, 7160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.1)
    add_table_source(doc, "Main Division Ranking System v1.1(2026.8.2) 요약 [I3].")

    doc.add_paragraph(style="Heading 2").add_run("5.2 사업적 역할")
    add_bullets(doc, [
        "Sub의 페이백 달성 이후에도 상위 성취 목표를 제공해 환급 직후 이탈을 줄인다.",
        "학습일수가 소진되면 Sub 강등과 재구독 경로가 열려 반복 결제 주기를 만든다.",
        "개인 상대 지정이 아닌 티어 선택·서버 무작위 배정으로 담합과 특정 상대 몰아주기를 줄인다.",
        "정책 버전, 원장, 서버 시간, 문제·답안·증거 기록으로 정산 결과를 재현할 수 있게 한다.",
    ])

    doc.add_paragraph(style="Heading 2").add_run("5.3 상용 운영 통제사항")
    risk_rows = [
        ("공정성", "문제 난이도·채점·상대 선정의 사후 조작 금지", "정책 버전 고정·서버 권위 판정"),
        ("과몰입", "학습일수 경쟁이 교육 목적을 압도할 위험", "일일 횟수 제한·보호자 안내·휴식 정책"),
        ("부정행위", "AI 대리풀이·공모·다계정·동일 파일", "풀이 증거·이상 징후·운영자 검토"),
        ("환급 분쟁", "조건 오해·정산 지연·지급 거절 분쟁", "결제 전 요약·근거 원장·이의신청 절차"),
        ("재무", "예상보다 높은 페이백률", "별도 충당금·월별 상한 경보·정책 신규 버전"),
    ]
    add_table(doc, ["위험", "내용", "통제"], risk_rows, [1500, 3750, 4110], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.9)

    new_page(doc)
    add_section_heading(doc, "06", "공동사업 운영 구조", "사업자등록용 공동사업과 실제 운영 책임을 일치시키는 내부 기준")
    doc.add_paragraph(style="Heading 2").add_run("6.1 공동사업 원칙")
    add_para(doc, "공동사업장은 하나의 사업자로 장부와 사업소득을 계산한 후, 동업계약에서 약정한 손익분배비율에 따라 각 공동사업자에게 소득을 배분한다. 사업자등록 시 공동사업자, 대표공동사업자, 손익분배비율, 지분·출자명세를 신고해야 하며, 형식상 비율과 실제 출자·업무·분배가 일치해야 한다. [S8]")

    doc.add_paragraph(style="Heading 2").add_run("6.2 권장 역할 분장")
    role_rows = [
        ("공동사업자 1", "제품·교육", "교육과정 설계, 콘텐츠 기준, 문제 품질, 학습정책, 사용자 리서치"),
        ("공동사업자 2", "기술·운영", "서비스 개발, 인프라, 데이터·보안, 결제·원장, 운영 자동화"),
        ("공동", "경영·통제", "가격·예산·채용·외주·정책 변경·분쟁·대외계약의 공동 의사결정"),
    ]
    add_table(doc, ["주체", "주관 영역", "주요 책임"], role_rows, [1900, 1800, 5660], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.1)

    doc.add_paragraph(style="Heading 2").add_run("6.3 동업계약서에 별도 확정할 항목")
    governance_rows = [
        ("출자·지분·손익", "현금·현물·기존 개발물·노동 기여의 평가, 지분·손익분배·적자 부담"),
        ("의사결정", "일상 집행 한도, 중요사항 전원동의, 교착상태 해결"),
        ("지식재산", "기존 코드·콘텐츠 귀속, 공동사업 중 산출물, 탈퇴 후 사용권"),
        ("보수·인출", "사업자 인출금, 비용 정산, 이익배당 시점, 세금 유보금"),
        ("계정·보안", "도메인·서버·DB·PG·은행·소스 저장소의 공동 통제"),
        ("탈퇴·분쟁·청산", "지분 평가·우선매수권·고객/데이터 인계 및 부채·환불·페이백 우선 변제"),
    ]
    add_table(doc, ["항목", "확정 내용"], governance_rows, [2200, 7160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.1)
    add_callout(doc, "청년 감면 유의", "공동사업에서 청년창업 대표자는 손익분배비율이 가장 큰 사업자다. 가장 큰 비율이 같은 사람이 둘 이상이면 그 모두가 청년 요건을 충족해야 한다. 따라서 50:50 구조에서 한 명만 청년이면 청년창업 감면이 배제될 수 있다. 비율은 절세만을 위해 형식적으로 정하지 않고 실제 출자·역할·분배와 일치시킨다. [S6][S8]", fill=CALLOUT, border=INK)

    new_page(doc)
    add_section_heading(doc, "07", "예상 매출·지출·이익", "완성된 상용 서비스의 12개월 예측 — 최초 결제자 1,000명, 가격 29,000원")
    doc.add_paragraph(style="Heading 2").add_run("7.1 산출 가정")
    assumption_rows = [
        ("분석대상", "최초 유료 결제자 1,000명 코호트"),
        ("가격", "결제 1건당 29,000원"),
        ("PG", "변동 수수료 3.74%, 연 고정비 330,000원 가정"),
        ("페이백", "매출 대비 평균 약 4.95%, 수령자 약 9.81%"),
        ("변동 운영비", "결제 1건당 5,000원(서버·문제·고객지원 통합 가정)"),
        ("추가 연간 예산", "마케팅·관리·보안·도구·예비비 13,000,000원"),
        ("제외", "공동사업자 본인 인출금·국민연금·건강보험·VAT 확정효과·소득세"),
    ]
    add_table(doc, ["항목", "가정"], assumption_rows, [2400, 6960], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.2)

    doc.add_paragraph(style="Heading 2").add_run("7.2 12개월 손익 시나리오")
    pnl_rows = [
        ("보수적", "1,447건", "41,956,040원", "24,213,474원", "17,742,566원", "42.3%"),
        ("기준", "1,940건", "56,261,450원", "27,946,297원", "28,315,153원", "50.3%"),
        ("강한 유지", "3,085건", "89,467,610원", "36,540,451원", "52,927,159원", "59.2%"),
    ]
    add_table(doc, ["시나리오", "연 결제", "고객 결제총액", "예상 지출", "세전 사업이익", "이익률"], pnl_rows,
              [1250, 1300, 1900, 1900, 1900, 1110],
              [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT], font_size=8.7)
    add_table_source(doc, "내부 시뮬레이션 [I4]에 결제 1건당 운영비 5,000원과 추가 연간 예산 13,000,000원을 반영. VAT·소득세 전.")

    doc.add_paragraph(style="Heading 2").add_run("7.3 기준 시나리오 지출 세부")
    expense_rows = [
        ("페이백", "2,811,869원", "조건 충족 이용자 환급"),
        ("PG 변동 수수료", "2,104,178원", "결제액의 3.74% 가정"),
        ("PG 고정비", "330,000원", "연간 내부 가정"),
        ("변동 운영비", "9,700,250원", "1,940건 × 약 5,000원"),
        ("마케팅·초기 유입", "6,000,000원", "월 500,000원 운영예산"),
        ("회계·세무·법률", "2,400,000원", "월평균 200,000원 상당"),
        ("보안·개인정보·보험", "1,200,000원", "연간 예비비"),
        ("개발 도구·콘텐츠 자산", "2,400,000원", "연간 운영예산"),
        ("일반 예비비", "1,000,000원", "분쟁·환급·장애 대응"),
        ("합계", "27,946,297원", "세전 사업비용"),
    ]
    table = add_table(doc, ["지출 항목", "연간 금액", "비고"], expense_rows, [3000, 2000, 4360], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.0)
    for cell in table.rows[-1].cells:
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    add_table_source(doc, "사업자 본인의 이익 인출은 공동사업 소득의 분배이며 본 표의 비용에 포함하지 않음.")

    doc.add_paragraph(style="Heading 2").add_run("7.4 규모 민감도")
    scale_rows = [
        ("500명", "약 970건", "28,130,725원", "기준 시나리오의 50% 선형 환산"),
        ("1,000명", "1,940건", "56,261,450원", "내부 기준 시나리오"),
        ("3,000명", "약 5,820건", "168,784,350원", "동일 재구독률·가격 가정"),
        ("5,000명", "약 9,700건", "281,307,250원", "동일 재구독률·가격 가정"),
    ]
    add_table(doc, ["최초 결제자", "예상 연 결제", "예상 고객 결제총액", "비고"], scale_rows, [1700, 1700, 2500, 3460], [WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.0)
    add_callout(doc, "손익분기 해석", "결제 1건당 직접 공헌이익을 약 21,150원으로 볼 때, 추가 연간 예산 1,300만 원을 충당하는 단순 손익분기점은 약 615건이다. 다만 대표자 생활비, 보험료, 세금, 환불·차지백, 신규 콘텐츠 확대비가 추가되면 실제 손익분기점은 높아진다.")

    new_page(doc)
    add_section_heading(doc, "08", "세무 및 절세 가능성", "2026년 8월 현재 법령 기준의 조건부 검토 — 확정 세무의견 아님")
    add_callout(doc, "가장 중요한 결론", "세액감면은 연매출을 바꾸지 않는다. 청년창업 감면은 소득세를 줄여 세후 유보금을 늘리고, 부가가치세 과·면세는 29,000원 소비자가격 중 얼마를 회계상 매출로 인식하는지를 바꾼다.", fill=CALLOUT, border=INK)

    doc.add_paragraph(style="Heading 2").add_run("8.1 청년창업중소기업 세액감면")
    add_para(doc, "2026년 1월 1일 이후 창업하는 청년창업중소기업은 최초 소득 발생연도부터 5년간 사업장 위치에 따라 소득세 감면을 받을 수 있다. 청년은 개인사업자 창업 당시 15세 이상 34세 이하이며, 병역 이행기간은 최대 6년까지 연령 계산에서 제외된다. 수도권과밀억제권역 50%, 그 밖의 수도권 75%, 수도권 밖 100%가 기본 감면율이다. [S5][S6]")
    youth_rows = [
        ("수도권과밀억제권역", "50%", "서울 등 해당 권역 여부를 주소별 확인"),
        ("수도권 내 비과밀 지역", "75%", "수도권 인구감소지역 포함 여부 별도 확인"),
        ("수도권 밖", "100%", "업종·창업 요건 충족 전제"),
    ]
    add_table(doc, ["2026년 이후 사업장 위치", "청년 기본 감면율", "유의사항"], youth_rows, [3200, 1800, 4360], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.1)
    add_table_source(doc, "자료: 국세청 청년 창업기업 세액감면 안내 및 조세특례제한법 제6조 [S5][S6].")

    doc.add_paragraph(style="Heading 3").add_run("업종 적격성")
    add_para(doc, "일반 고등수학 교육서비스는 조세특례제한법 제6조 제3항의 열거 업종에 포함되지 않는다. 교육 분야 중 명시적으로 포함되는 것은 직업기술 분야 학원 또는 직업능력개발훈련시설 등으로 제한된다. 반면 정보통신업은 열거 업종이므로 Matths의 실제 주된 사업이 소프트웨어·정보서비스인지, 교육용역인지에 따라 결론이 달라질 수 있다. 업종코드만 정보통신업으로 선택하는 것은 충분하지 않으며 계약, 화면, 매출 발생 원인, 콘텐츠 제공 방식 등 사업의 실질이 일치해야 한다. [S6]")
    add_callout(doc, "보수적 판단", "현재 사업 설명이 ‘고등학생에게 수학을 가르치는 유료 서비스’에 가깝다면 청년창업 감면 0%를 기본값으로 두고, 사업자등록 전 국세청의 소득세 공제·감면 컨설팅 또는 세무대리인 서면검토를 받아 정보통신업 해당 가능성을 확인하는 것이 안전하다. [S7]", fill=CALLOUT, border=INK)

    doc.add_paragraph(style="Heading 2").add_run("8.2 부가가치세 과·면세")
    add_para(doc, "교육 목적만으로 부가가치세가 면제되지는 않는다. 면세 교육용역은 원칙적으로 주무관청의 허가·인가를 받거나 등록·신고된 학교·학원·강습소·교습소 등이 학생에게 지식·기술을 가르치는 경우다. 온라인 플랫폼이 이러한 지위를 갖추지 못하면 과세사업으로 판단될 가능성을 먼저 검토해야 한다. [S9]")
    vat_rows = [
        ("교육용역 면세 인정", "56,261,450원", "28,315,153원", "매입세액 불공제 가능성 반영 전"),
        ("VAT 과세·29,000원 포함", "51,146,773원", "23,200,476원", "산출 VAT 5,114,677원 단순 분리, 매입세액공제 전"),
    ]
    add_table(doc, ["기준 시나리오", "회계상 매출", "잠정 세전이익", "비고"], vat_rows, [2800, 2100, 2100, 2360], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.8)
    add_table_source(doc, "고객 결제총액 56,261,450원을 기준으로 단순 비교. 실제 납부 VAT는 매입세액·간이과세·면세 겸영 여부에 따라 달라짐.")

    doc.add_paragraph(style="Heading 2").add_run("8.3 세액감면에 따른 세후 유보금 변화")
    add_para(doc, "아래는 ‘VAT 과세, 기준 시나리오 잠정 세전이익 23,200,476원, 공동사업자 2인 50:50, 두 사람 모두 청년, 다른 종합소득 없음, 본인 기본공제 각 150만 원’이라는 제한된 예시다. 공동사업소득은 1인당 약 1,160만 원, 과세표준은 약 1,010만 원으로 보고 6% 세율을 적용했다. 지방소득세는 보수적으로 국세 산출세액의 약 10% 수준을 별도 유지해 표시했다. [S10][S11]")
    tax_rows = [
        ("감면 없음", "56,261,450원", "1,212,029원", "121,203원", "21,867,244원", "-"),
        ("청년 50%", "56,261,450원", "606,014원", "121,203원", "22,473,259원", "+606,014원"),
        ("청년 75%", "56,261,450원", "303,007원", "121,203원", "22,776,265원", "+909,021원"),
        ("청년 100%", "56,261,450원", "0원", "121,203원", "23,079,273원", "+1,212,029원"),
    ]
    add_table(doc, ["국세 감면", "고객 결제총액", "예상 소득세", "지방세 가정", "세후 유보금", "절세효과"], tax_rows,
              [1400, 1800, 1600, 1450, 1650, 1460],
              [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT], font_size=8.4)
    add_table_source(doc, "단순 예시이며 건강보험료·연금·타소득·세액공제·최저한세·농어촌특별세·지방세 감면·매입 VAT를 반영하지 않음. 실제 신고액이 아님.")

    hard_new_page(doc)
    doc.add_paragraph(style="Heading 2").add_run("8.4 합법적 절세 실행항목")
    add_para(doc, "절세는 매출을 인위적으로 변경하는 수단이 아니라, 적격 업종·증빙·공동사업 정합성·정확한 신고를 통해 세후 유보금을 합법적으로 확보하는 관리절차다.")
    add_numbered(doc, [
        "사업자등록 전 실제 주업종, 사업장 소재지, 창업 해당 여부, 공동사업 청년 대표 요건을 서면 검토한다.",
        "사업용 계좌·카드·PG 정산계정을 분리하고 서버·콘텐츠·광고·외주·법률·회계비용의 적격증빙을 수취한다.",
        "페이백 예정액과 환불·차지백을 월별 충당금으로 관리하고 실제 지급 증빙을 보관한다.",
        "공동사업자 손익분배비율을 동업계약·사업자등록·장부·실제 분배에 동일하게 반영한다.",
        "노란우산공제, 통합투자세액공제, 고용 관련 공제는 실제 가입·투자·채용 시 별도 검토한다.",
        "정식 결제 운영 시 부가가치세 과·면세와 소비자가격 표시 방식을 확정하고, 면세를 전제로 가격을 정하지 않는다.",
    ])

    new_page(doc)
    add_section_heading(doc, "09", "운영·성장계획 및 성과관리", "개발 완료 이후 시장 진입과 학생·구매자별 성과관리 계획")
    doc.add_paragraph(style="Heading 2").add_run("9.1 단계별 실행")
    roadmap_rows = [
        ("1단계", "사업 운영 개시", "0~1개월", "동업계약, 업종·VAT 검토, 사업자등록, 계좌·장부, PG·약관·개인정보 확정"),
        ("2단계", "초기 상용화", "1~3개월", "학생 체험·가입 유입, 학부모 결제전환, 제3자 학습권 구매 절차 운영"),
        ("3단계", "유료고객 안정화", "3~6개월", "100~500명 기준 가격·페이백·재구매·환급·CS 지표 검증"),
        ("4단계", "기준 코호트 확보", "6~12개월", "최초 결제자 1,000명, 구매자·학생 코호트 분리분석, 손익 재산출"),
        ("5단계", "채널 확장", "12개월 이후", "학교·장학·후원 등 제3자 구매채널과 B2B·B2G 제휴 검토"),
    ]
    add_table(doc, ["단계", "목표", "기간", "완료 기준"], roadmap_rows, [1100, 1800, 1300, 5160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.9)

    doc.add_paragraph(style="Heading 2").add_run("9.2 핵심 KPI")
    kpi_rows = [
        ("학생 획득", "방문→가입, 체험→학습개시, 활성 학생, 추천 유입", "주·월"),
        ("구매자·학습권", "결제전환·구매자 유형·갱신, 학생 지정·사용개시·미사용권", "주·월"),
        ("학습·평가", "활성일, 완료율, 오답 재도전, 주간 평가·점수 변화", "일·주"),
        ("수익·페이백", "결제·재구매·CAC·공헌이익, 수령률·지급액/매출·이의신청", "월"),
        ("Main·품질", "진입·체류·강등·72시간 재구독, 오류·환불·부정행위·개인정보", "월·상시"),
    ]
    add_table(doc, ["영역", "핵심 지표", "주기"], kpi_rows, [1600, 6300, 1460], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER], font_size=8.8)

    doc.add_paragraph(style="Heading 2").add_run("9.3 모델 재산출 기준")
    add_para(doc, "구매자 유형별 결제·재구매, 학생별 학습개시·유지, 구매자-학생 연결, Sub 만료, Main 강등, 페이백 심사 및 경기 결과가 각 100건 이상 축적되면 전환율·결제주기·학생 체류기간·페이백률을 다시 추정한다. 가정과 실제 값의 차이가 20%를 넘거나 지급액/매출이 사전 경보선을 넘으면 신규 정책 버전에서 가격·조건·원가를 조정한다. 이미 시작된 주기에는 소급 적용하지 않는다.")

    new_page(doc)
    add_section_heading(doc, "10", "주요 위험 및 대응", "사업·재무·규제·운영 위험을 사전에 분리 관리")
    risk_full_rows = [
        ("수요", "시각화 콘텐츠의 유료 전환이 낮음", "중", "과목별 무료 체험, 학생 유입·구매자 전환 퍼널 분리 측정"),
        ("비대칭 고객", "학생은 원하지만 구매자가 결제하지 않거나, 구매 후 학생이 미사용", "상", "학생·구매자별 안내, 학습권 지정, 미사용 알림, 갱신 근거 제공"),
        ("리텐션", "랭킹·페이백이 학습보다 경쟁만 자극", "중", "학습 완료를 필수 조건으로 연결, 보호자·청소년 정책"),
        ("페이백", "실제 지급률이 5% 가정을 상회", "상", "월별 충당금, 경보선, 신규 정책 버전, 과장광고 금지"),
        ("VAT", "면세 오판으로 세금·가산세 발생", "상", "등록 전 서면검토, 과세 기준 가격 확보"),
        ("청년감면", "업종·지분·지역 요건 미충족", "중", "감면 0%를 기본 예산으로 보고 사전 컨설팅"),
        ("공동사업", "역할·지분·IP·탈퇴 분쟁", "상", "상세 동업계약, 계정 공동통제, 중요사항 전원동의"),
        ("개인정보", "미성년자·답안·결제·증거자료 유출", "상", "최소수집, 분리보관, 접근통제, 보존기간, 사고대응"),
        ("공정성", "AI 대리풀이·공모·다계정", "중", "서버 판정, 증거 제출, 이상징후, 보류·이의신청"),
        ("콘텐츠", "오답·교육과정 불일치·저작권", "상", "출처관리, 이중검수, 자체 생성 문제, 정정 이력"),
        ("인프라", "시험·정산 시 장애", "중", "동결·재시도·멱등성·백업·장애 공지"),
    ]
    add_table(doc, ["영역", "위험", "등급", "대응"], risk_full_rows, [1400, 3150, 850, 3960], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.55)

    doc.add_paragraph(style="Heading 2").add_run("승인 전 필수 체크")
    add_bullets(doc, [
        "공동사업 동업계약, 손익분배비율·대표공동사업자, 주업종·VAT·청년감면 검토를 확정한다.",
        "PG·소비자가격·페이백 충당금·지급·이의신청·환불 절차를 동일한 정책으로 연결한다.",
        "미성년자 개인정보·풀이 증거·결제정보의 수집·보관·파기와 약관·결제 전 요약을 검수한다.",
        "상용 운영의 장애·보안·CS 담당자와 비상 연락체계를 지정한다.",
    ])

    new_page(doc)
    add_section_heading(doc, "11", "종합 결론 및 공동사업자 확인", "본 문서의 가정과 조건을 전제로 한 내부 사업 추진 판단")
    doc.add_paragraph(style="Heading 2").add_run("11.1 종합 결론")
    add_para(doc, "Matths는 2022 개정 고등학교 수학 교육과정 13개 과목 전체를 제공하는 완성형 시각화 학습 플랫폼이며, 고등학생 사교육 참여자의 월평균 지출 79만 3천 원 대비 약 3.7%의 가격으로 진입할 수 있다. 내부 기준 시나리오에서는 최초 결제자 1,000명으로 연간 약 1,940건의 결제와 고객 결제총액 약 5,626만 원을 예상한다. 페이백·PG·결제건당 운영비·추가 연간 예산을 반영한 세전 사업이익은 VAT 판단 전 기준 약 2,832만 원이다.")
    add_para(doc, "핵심 사업구조는 학부모·보호자·제3자가 학습권을 구매하고 학생이 실제 서비스를 이용하는 비대칭 시장에 있다. 따라서 학생의 학습효과·동기·편의와 구매자의 가격·신뢰·성과확인을 동시에 만족시켜야 한다. 사업성 판단의 주요 변수는 ① 구매자 유형별 결제·재구매, ② 학생의 실제 사용·유지, ③ 페이백 지급률과 Main 체류효과, ④ 부가가치세 과·면세, ⑤ 청년창업 세액감면 대상 업종 인정 여부다.")
    add_callout(doc, "권고", "완성된 제품의 상용 운영을 전제로 사업자등록을 추진하되, ‘공동사업 동업계약 + 업종·VAT·청년감면 서면검토 + 페이백 충당금·분쟁절차 + 구매자·학생 권한 및 개인정보 기준’의 네 가지 행정·통제조건을 확정한다.")

    doc.add_paragraph(style="Heading 2").add_run("11.2 공동사업자 확인")
    add_para(doc, "아래 공동사업자는 본 문서가 제품 개발 완료와 상용 운영 가능 상태를 전제로 하되, 매출·비용·세무 수치에는 실제 유료 운영 실적이 아닌 예측 가정이 포함되어 있음을 확인한다. 사업자등록 신청서와 동업계약서에서 확정한 사항은 본 문서보다 우선한다.")
    sign_rows = [
        ("공동사업자 1", "성명: ____________________", "서명/인: ____________________", "일자: 20____. ____. ____."),
        ("공동사업자 2", "성명: ____________________", "서명/인: ____________________", "일자: 20____. ____. ____."),
        ("대표공동사업자", "성명: ____________________", "서명/인: ____________________", "일자: 20____. ____. ____."),
    ]
    add_table(doc, ["구분", "성명", "확인", "일자"], sign_rows, [1700, 2600, 2600, 2460], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.1)

    new_page(doc)
    add_section_heading(doc, "A", "근거자료 및 출처", "외부 공식자료와 내부 정책·시뮬레이션을 구분")
    doc.add_paragraph(style="Heading 2").add_run("A.1 외부 공식자료")
    add_source(doc, "S1", "2025년 초중고사교육비조사 결과", "국가데이터처(통계청)·교육부", "2026.03.12", "https://sri.kostat.go.kr/boardDownload.es?bid=245&list_no=443953&seq=1", "고등학교 참여학생 월평균 79.3만 원, 고1 80.6만 원")
    add_source(doc, "S2", "2025년 초·중·고 사교육비 조사결과 카드뉴스", "교육부", "2026.05.12", "https://www.moe.go.kr/boardCnts/viewRenew.do?boardID=340&boardSeq=106117&lev=0&m=0206&opType=N&page=3&s=moe", "고등학교 사교육비 7.8조 원, 참여율 63.0%")
    add_source(doc, "S3", "2025-4 예비창업패키지 사업실명제", "창업진흥원", "2025", "https://www.kised.or.kr/board.es?act=view&bid=0002&list_no=4089&mid=a10106000000", "공식 창업지원 사업계획서 참고")
    add_source(doc, "S4", "사업계획서 주요 작성항목", "중소벤처기업부", "2023", "https://www.mss.go.kr/common/board/Download.do?bcIdx=1041279&cbIdx=86&streFileNm=f2f83f71-94cb-49d5-9d3d-0f2445e65c26.pdf", "필요성·경쟁사·확대전략·자금운용·팀 구성")
    add_source(doc, "S5", "청년 창업기업에 대한 세액감면 적용 사전 안내", "국세청", "2026", "https://s.nts.go.kr/nts/cm/cntnts/cntntsView.do?cntntsId=239082&mi=41127", "2026년 이후 지역별 50%·75%·100%, 5년")
    add_source(doc, "S6", "조세특례제한법 제6조 및 시행령 제5조", "국가법령정보센터", "2026.07.01 시행", "https://www.law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1032222983", "감면 대상 업종·지역별 비율·청년 요건")
    add_source(doc, "S7", "소득세 공제·감면 컨설팅 제도", "국세청", "2026", "https://s.nts.go.kr/nts/cm/cntnts/cntntsView.do?cntntsId=239070&mi=41093", "개인 중소기업 공제·감면 사전검토")
    add_source(doc, "S8", "소득세법 제43조 공동사업 소득금액 계산", "국가법령정보센터", "2026.07.01 시행", "https://law.go.kr/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1021859115", "공동사업장 소득을 약정 손익분배비율로 배분")
    add_source(doc, "S9", "부가가치세법 시행령 제36조 면세 교육용역 범위", "국가법령정보센터", "2026.04.01 시행", "https://law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lspttninfSeq=112803", "주무관청 허가·인가·등록·신고 시설 등")
    add_source(doc, "S10", "소득세법 제55조 종합소득세율", "국가법령정보센터", "2026.07.01 시행", "https://law.go.kr/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1023583825", "과세표준 1,400만 원 이하 6%")
    add_source(doc, "S11", "소득세법 제50조 기본공제", "국가법령정보센터", "2026.07.01 시행", "https://www.law.go.kr/lsLinkCommonInfo.do?lsJoLnkSeq=1032954559", "본인 등 1명당 연 150만 원")
    add_source(doc, "S12", "교육부 고시 제2022-33호 [별책 8] 수학과 교육과정", "교육부", "2022", "https://www.moe.go.kr/boardCnts/viewRenew.do?boardID=141&boardSeq=93458&lev=0", "2022 개정 고등학교 수학 교육과정의 과목 편성 근거")

    doc.add_paragraph(style="Heading 2").add_run("A.2 내부자료")
    internal_sources = [
        ("I1", "curriculum_folder/*.yaml 및 서비스 콘텐츠 카탈로그", "고등학교 수학 13개 과목 구성·학습 콘텐츠 범위"),
        ("I2", "docs/logic/01_MATTHS_CURRENT_SYSTEM.md", "현행 기술·학습·평가·운영 구조"),
        ("I3", "docs/logic/03_SUB_DIVISION_RANKING_SYSTEM_PAYBACK.md 및 04_MAIN_DIVISION_RANKING_SYSTEM.md", "Sub 페이백과 Main Division 운영 규칙"),
        ("I4", "docs/logic/09_GOAT_ARENA_PROFIT_LOSS_SIMULATION.md", "1,000명 코호트·결제·페이백·PG·재구독 손익 추정"),
        ("I5", "output/legal/Matths_서비스_적법성_검토_및_법률적_소명서_v1.0.md", "교육서비스·내부 일수·페이백 구조에 관한 내부 검토"),
    ]
    for code, path, note in internal_sources:
        p = doc.add_paragraph(style="Source List")
        set_font(p.add_run(f"[{code}] {path}"), size=8.1, bold=True, color=TEXT)
        set_font(p.add_run(f" — {note}"), size=8.1, color=MUTED)

    doc.add_paragraph(style="Heading 2").add_run("A.3 산출 및 해석 유의사항")
    add_bullets(doc, [
        "제품 개발은 완료된 것으로 전제했으며, 시장 통계는 2025년 조사 결과를 사용하고 재무 수치는 실제 매출 실적이 아닌 내부 정책 검증 모델로 산출했다.",
        "세무 법령은 2026년 8월 2일 시행 기준이며, ‘연매출’은 고객 결제총액과 회계상 공급가액을 구분했다. VAT 과세 시 소비자가격에 포함된 VAT는 매출이 아니다.",
        "법률·세무 적용은 사실관계에 따라 달라지므로 관할 세무서, 국세청 사전컨설팅, 세무대리인·변호사의 검토로 확정한다.",
    ])

    doc.add_paragraph(style="Heading 2").add_run("A.4 문서 통제")
    control_rows = [
        ("v1.0", "2026.08.02", "최초 작성", "시장·수익·Main Division·공동사업·세무 시나리오 통합"),
        ("v1.1", "2026.08.02", "전면 개정", "13개 과목 전체·개발 완료 전제·비대칭 시장·맑은 고딕·페이지 편집 반영"),
    ]
    add_table(doc, ["버전", "기준일", "구분", "변경 내용"], control_rows, [1200, 1600, 1500, 5060], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9.0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
