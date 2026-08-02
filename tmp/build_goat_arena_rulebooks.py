from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/sangyoonlee/Desktop/SangYoon Lee/SINGAPORE 2025-/Personal Projects/Matths")
OUT_DIR = ROOT / "output" / "rulebooks"
FONT = "Malgun Gothic"

# compact_reference_guide, with a named Malgun Gothic + restrained user-guide color override.
INK = "182026"
NAVY = "233B4D"
BLUE = "365F78"
MUTED = "69757D"
LIGHT = "F1F4F6"
PALE = "E8EEF2"
CALLOUT = "F5F7F8"
BORDER = "B7C1C7"
WHITE = "FFFFFF"

PAGE_DXA = 9360
TABLE_INDENT = 120


def set_font(run, size=None, bold=None, italic=None, color=INK):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], indent=TABLE_INDENT):
    if sum(widths) != PAGE_DXA:
        raise ValueError(f"Table widths must total {PAGE_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BLUE, size=12, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "6")
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_font(run, size=8.5, color=MUTED)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id, fmt, marker, left=540, hanging=270):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)

    add_abstract(950, "bullet", "•")
    add_abstract(951, "decimal", "%1.")
    add_num(950, 950)
    add_num(951, 951)


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name, size, color, bold in (
        ("Kicker", 8.7, BLUE, True),
        ("Lead", 11.2, INK, False),
        ("Small Note", 8.7, MUTED, False),
    ):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.2

    styles["Lead"].paragraph_format.space_after = Pt(10)
    styles["Lead"].paragraph_format.line_spacing = 1.22


def setup_document(book_title: str, short_title: str, version: str):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False

    configure_styles(doc)
    configure_numbering(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_font(hp.add_run("MATTHS  |  GOAT ARENA 사용자 룰북"), size=8.2, bold=True, color=MUTED)
    set_font(hp.add_run(f"\t{short_title}"), size=8.2, bold=True, color=NAVY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(fp.add_run(f"{version}  |  "), size=8.2, color=MUTED)
    add_page_field(fp)

    doc.core_properties.title = book_title
    doc.core_properties.subject = "GOAT Arena 사용자용 확정 규칙"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "Matths, GOAT Arena, 사용자 룰북"
    doc.core_properties.comments = ""
    return doc


def add_cover(doc, category: str, title: str, subtitle: str, policy: str, chapters: Sequence[str]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(12)
    set_font(p.add_run("MATTHS  ·  GOAT ARENA"), size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(category), size=12, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_together = True
    set_font(p.add_run(title), size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    set_font(p.add_run(subtitle), size=12.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(policy), size=9.2, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    set_font(p.add_run("적용 시간대  ·  대한민국 표준시(KST)"), size=8.8, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run("이 룰북에서 확인할 내용"), size=10.5, bold=True, color=NAVY)
    for chapter in chapters:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(2)
        apply_num(p, 950)
        set_font(p.add_run(chapter), size=9.2, color=INK)


def chapter(doc, number: int, title: str, lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    set_font(p.add_run(""), size=1, color=WHITE)

    k = doc.add_paragraph(style="Kicker")
    k.paragraph_format.keep_with_next = True
    set_font(k.add_run(f"CHAPTER {number:02d}"), size=8.7, bold=True, color=BLUE)
    h = doc.add_paragraph(style="Heading 1")
    set_font(h.add_run(title), size=16, bold=True, color=NAVY)
    if lead:
        p = doc.add_paragraph(style="Lead")
        set_font(p.add_run(lead), size=11.2, color=INK)


def heading(doc, text: str, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_font(p.add_run(text), size=13 if level == 2 else 11.5, bold=True, color=BLUE if level == 2 else NAVY)
    return p


def para(doc, text: str, bold_prefix: str | None = None, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.18
        apply_num(p, 950)
        set_font(p.add_run(item))


def numbered(doc, items: Iterable[str]):
    next_id = getattr(doc, "_rulebook_next_num_id", 952)
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), "951")
    num.append(ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    doc._rulebook_next_num_id = next_id + 1
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.18
        apply_num(p, next_id)
        set_font(p.add_run(item))


def callout(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, CALLOUT)
    set_paragraph_border(p)
    set_font(p.add_run(label + "  "), size=9.8, bold=True, color=BLUE)
    set_font(p.add_run(text), size=9.8, color=INK)
    return p


def table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], font_size=9.1):
    tbl = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(tbl, widths)
    set_table_borders(tbl)
    repeat_table_header(tbl.rows[0])
    for idx, value in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(value), size=9.2, bold=True, color=NAVY)
    for values in rows:
        row = tbl.add_row()
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.10
            set_font(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(tbl, widths)
    for row in tbl.rows:
        prevent_row_split(row)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    set_font(spacer.add_run(""), size=2, color=WHITE)
    return tbl


def glossary(doc, rows):
    table(doc, ["용어", "뜻"], rows, [2300, 7060], font_size=9.3)


def save(doc, filename):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_common():
    doc = setup_document("GOAT Arena 공통 규칙 사용자 룰북", "공통 규칙", "사용자용 v1.0")
    add_cover(
        doc,
        "공통 규칙",
        "GOAT Arena\n사용자 룰북",
        "Sub와 Main에서 함께 적용되는 경기·판정·보호 기준",
        "사용자용 v1.0  |  현재 확정 공통 규칙",
        [
            "Division과 랭킹의 구분",
            "공식 경기 참가 자격",
            "경기 신청부터 정산까지의 흐름",
            "문제 풀이와 증거 제출",
            "승패와 Arena 상태 변경",
            "복수전·미완료·장애 처리",
            "개인정보·공정성·일요일 잠금",
        ],
    )

    chapter(doc, 1, "먼저 알아둘 구조", "공통 규칙은 Sub Division과 Main Division의 모든 공식 1대1 경기에 함께 적용됩니다. 학습일수의 예치·이전·소각과 Division별 상대 조건은 각 Division 룰북을 따릅니다.")
    glossary(doc, [
        ("Sub Division", "결제주기별 페이백과 Main 진입 자격을 함께 경쟁하는 Division"),
        ("Main Division", "Sub에서 진입 자격을 얻은 사용자가 학습일수를 예치하고 경쟁하는 상위 Division"),
        ("Arena 상태", "Arena 랭크, 랭크 안의 정확한 순위, Arena GP를 합친 경기 상태"),
        ("Skill MMR", "배치고사와 매주 일요일 공식 모의고사로만 바뀌는 시험 실력 지표"),
        ("Final Ranking", "Skill MMR과 현재 Division 성과 등을 함께 반영하는 전체 종합 랭킹"),
        ("도전자", "상대의 Arena 상태에 도전하는 사용자"),
        ("방어자", "현재 Arena 상태를 지키는 사용자"),
    ])
    heading(doc, "Sub와 Main은 서로 분리됩니다")
    bullets(doc, [
        "서로 다른 Division 사용자는 공식 랭킹 경기에서 만나지 않습니다.",
        "Sub와 Main은 참가자 풀, Arena 랭킹, 학습일수 운영 방식과 경기 정책을 각각 따로 운영합니다.",
        "Sub에서 얻은 결과를 Main 경기 결과로 임의 전환하지 않습니다.",
        "Main 진입 자격을 얻었다는 사실만으로 Sub 경기 규칙이 Main에 그대로 적용되지는 않습니다.",
    ])
    callout(doc, "핵심 구분", "1대1 경기는 Arena 상태를 다루며 Skill MMR을 바꾸지 않습니다. Skill MMR은 배치고사와 매주 일요일 공식 모의고사에서만 바뀝니다.")

    chapter(doc, 2, "공식 경기 참가 자격", "공식 경기에는 활성 계정과 유효한 유료 학습권을 가진 사용자만 참여할 수 있습니다.")
    heading(doc, "참가 전에 모두 충족해야 하는 조건")
    bullets(doc, [
        "현재 시즌의 필요한 배치 절차를 완료해야 합니다.",
        "계정이 정상 이용 상태여야 합니다.",
        "유료 학습권이 활성 상태여야 합니다.",
        "사용 가능한 학습일수가 남아 있어야 합니다.",
        "일요일 Division 잠금 시간이 아니어야 합니다.",
        "정산이 끝나지 않은 다른 공식 경기가 없어야 합니다.",
    ])
    heading(doc, "모의고사 전용 패키지")
    para(doc, "모의고사 전용 패키지는 주간 공식 모의고사만 이용할 수 있습니다. 배치고사와 GOAT Arena 일반 쟁탈전·복수전에는 참가할 수 없으며, 이 패키지로 얻은 MMR만으로 Arena 입장 자격이나 Arena 상태를 받을 수 없습니다.")
    heading(doc, "학습일수가 끝난 경우")
    bullets(doc, [
        "Sub 사용자는 공격·방어·복수전과 신규 매칭이 즉시 중단됩니다.",
        "Main 사용자는 예약 중이거나 경기에 예치된 학습일수와 미정산 경기를 먼저 처리한 뒤 최종 만료 여부를 확정합니다.",
        "만료 사용자는 방어 후보로 남지 않습니다.",
        "진행 중 경기에 예치된 학습일수가 있으면 경기 정산을 먼저 진행합니다.",
    ])

    chapter(doc, 3, "경기가 진행되는 순서", "두 참가자는 동시에 접속할 필요가 없지만, 각자 경기 시작 뒤에는 개인 제한 시간이 계속 흐릅니다.")
    numbered(doc, [
        "서버가 계정·구독·학습일수·Division 참가 자격을 확인합니다.",
        "Division 규칙에 따라 상대 후보를 정하고 다시 자격을 확인합니다.",
        "경기 조건과 문제·채점 기준을 고정합니다.",
        "자동 검산을 통과한 문제를 확정하고 필요한 학습일수를 예약하거나 잠급니다.",
        "각 참가자가 자신의 경기 화면에서 문제 풀이를 시작합니다.",
        "답안과 풀이 활동을 저장하고 제한 시간 안에 제출합니다.",
        "풀이 증거를 제출하면 서버가 채점과 공정성 검사를 진행합니다.",
        "결과가 확정되면 Arena 상태와 학습일수를 한 번에 정산합니다.",
    ])
    heading(doc, "동시에 가질 수 있는 경기")
    para(doc, "한 사용자는 같은 시점에 정산되지 않은 공식 경기를 하나만 가질 수 있습니다. 결과가 검토 대기 상태이면 Arena 상태와 학습일수는 확정되지 않습니다.")
    heading(doc, "경기 조건은 중간에 바뀌지 않습니다")
    para(doc, "경기가 시작될 때 정해진 문제, 배점, 제한 시간, 채점 기준과 Division 규칙은 경기 종료까지 그대로 적용됩니다.")

    chapter(doc, 4, "문제 풀이와 제출", "모든 공식 1대1 경기는 두 참가자에게 같은 문제와 같은 제한 시간을 제공합니다.")
    table(doc, ["항목", "공통 기준"], [
        ("문제 형식", "주관식 준킬러 5문항"),
        ("총점", "100점"),
        ("제한 시간", "10분"),
        ("문제 조건", "같은 교육과정 범위·문항·배점·목표 난이도"),
        ("풀이 증거", "문제 종료 뒤 60초 안에 사진 1~5장 제출"),
    ], [2300, 7060], font_size=9.4)
    heading(doc, "문제 이동")
    bullets(doc, [
        "화면에는 현재 풀고 있는 한 문항만 표시됩니다.",
        "현재 문항은 다음 문제로 이동하기 전까지만 답을 바꿀 수 있습니다.",
        "다음 문제를 누르면 이전 문항의 답이 확정되며, 이전 문제와 답안을 다시 보거나 수정할 수 없습니다.",
        "5번 문항에서 다음 문제를 누르거나 10분이 지나면 문제 화면이 닫히고 풀이 증거 제출 화면으로 전환됩니다.",
    ])
    heading(doc, "시간과 자동 제출")
    bullets(doc, [
        "각 사용자가 시작한 서버 시각부터 개인 제한 시간이 흐릅니다.",
        "화면을 닫거나 다른 화면으로 이동해도 제한 시간은 멈추지 않습니다.",
        "제한 시간이 끝날 때 제출하지 않은 경우 마지막으로 저장된 답안이 자동 제출됩니다.",
        "분수·소수·동치식은 수학적으로 같은 값이면 정답으로 인정합니다.",
        "풀이 증거는 상대에게 공개되지 않으며 운영자는 공정성 확인을 위해 확인할 수 있습니다.",
    ])

    chapter(doc, 5, "승패와 Arena 상태", "점수와 정답 성과가 풀이시간보다 먼저 적용됩니다.")
    heading(doc, "승패 판정 순서")
    numbered(doc, [
        "총점이 높은 사용자가 승리합니다.",
        "총점이 같으면 정답 수가 많은 사용자가 승리합니다.",
        "정답 수도 같으면 정답을 맞힌 문항의 풀이시간이 짧은 사용자가 승리합니다.",
        "여전히 같으면 전체 풀이시간이 짧은 사용자가 승리합니다.",
        "모든 기준이 같으면 방어자가 승리합니다.",
    ])
    callout(doc, "추가 문제 없음", "현재 규칙에는 승부를 다시 가르는 Sudden Death가 없습니다. 모든 기준이 같을 때는 방어자 승리로 확정합니다.")
    heading(doc, "Arena 상태 변경")
    table(doc, ["경기 결과", "Arena 상태 처리", "Skill MMR"], [
        ("도전자 승리", "두 사용자의 Arena 랭크·정확한 순위·GP를 전부 교환", "변경 없음"),
        ("방어자 승리", "Arena 상태 교환 없음", "변경 없음"),
    ], [2100, 5260, 2000], font_size=9.2)
    heading(doc, "Final Ranking과의 관계")
    bullets(doc, [
        "Final Rating이나 Final Rank는 상대 선정, 문제 난이도, 승패 판정에 사용하지 않습니다.",
        "1대1 경기 결과만을 별도 보너스로 계속 누적하지 않습니다.",
        "Arena 상태가 정상 정산되면 확정된 정보를 바탕으로 Final Ranking을 다시 계산합니다.",
        "Final Ranking 계산 결과가 Skill MMR이나 Arena 상태를 거꾸로 바꾸지는 않습니다.",
    ])

    chapter(doc, 6, "복수전과 미완료 처리", "복수전의 경기 기준은 일반 경기와 같고 학습일수 정산은 각 Division 규칙을 따릅니다.")
    heading(doc, "복수전 선택권")
    bullets(doc, [
        "가장 최근에 끝난 원경기의 패자에게 결과 화면에서 한 번만 복수전 선택권이 주어지며, 선택하면 상대는 자동 참가합니다.",
        "경기 종료를 선택하거나 복수전 권리를 포기하면 같은 원경기로 다시 신청할 수 없습니다.",
        "신청 뒤 24시간 안에 양측이 문제 풀이를 완료해야 하며, 완료 기한은 일요일 14:30을 넘길 수 없습니다.",
    ])
    heading(doc, "미시작·미완료")
    bullets(doc, [
        "일반 쟁탈전은 공식 매치가 성립한 뒤 24시간 안에 시작해야 합니다.",
        "복수전 신청자가 기한 안에 완료하지 않으면 방어자 승리로 처리하고 Arena 상태는 교환하지 않습니다.",
        "학습일수 반환·이전·소각은 Sub 또는 Main 룰북의 복수전 정산 기준을 적용합니다.",
        "Main의 상위 티어 초대는 매치가 성립되기 전까지 고정 24시간 만료를 적용하지 않습니다.",
    ])
    heading(doc, "오류와 장애")
    table(doc, ["상황", "처리"], [
        ("적격 상대가 없음", "학습일수를 잠그지 않고 신청 종료"),
        ("서버 장애", "경기 무효 처리 후 예치 학습일수 반환"),
        ("문제 오류", "검토 대기로 전환한 뒤 재채점하거나 무효 처리"),
        ("결과가 확정되지 않음", "관련 정산과 페이백 심사를 보류"),
        ("정지·탈퇴·결제 분쟁", "발생 원인과 시점에 따라 검토 대기 또는 무효 처리"),
    ], [2650, 6710], font_size=9.2)

    chapter(doc, 7, "개인정보·공정성·일요일 잠금", "상대 선정과 판정은 서버 기준이며 경기 화면에는 필요한 최소 정보만 공개합니다.")
    heading(doc, "상대에게 공개되는 정보")
    bullets(doc, [
        "상대가 선정되기 전에는 개인 후보와 닉네임을 공개하지 않습니다.",
        "매치가 성립한 뒤에는 서비스 닉네임과 경기에 필요한 Arena 정보만 표시합니다.",
        "실명·학교·지역·연락처·결제 정보는 공개하지 않으며, 탈퇴 사용자의 과거 기록은 공개 화면에서 익명화합니다.",
    ])
    heading(doc, "부정행위와 공모")
    bullets(doc, [
        "연관 계정끼리의 반복 경기와 한 방향으로만 이어지는 학습일수 이전을 확인합니다.",
        "빠른 제출·반복 미응답·동일 답안·고의 오답·반복 미완료·화면 정보 조작을 확인합니다.",
        "의심 경기는 검토 대기로 전환되며, 검토가 끝날 때까지 Arena 상태와 학습일수를 확정하지 않습니다.",
    ])
    heading(doc, "매주 일요일 운영 시간")
    table(doc, ["시간", "적용 규칙"], [
        ("14:30부터", "Sub·Main 신규 경기 신청·수락·준비·시작 차단"),
        ("15:00부터", "Arena 상태를 바꾸는 경기 기록과 정산 잠금, 공개 Final Ranking 고정"),
        ("15:00~24:00", "정상 이용자의 주간 공식 모의고사와 새 랭킹 반영 준비"),
        ("월요일 00:00", "새 MMR·주간 보너스·Final Rank 공개 및 Arena 잠금 해제"),
    ], [2300, 7060], font_size=9.2)
    para(doc, "14:30 전에 시작한 경기는 15:00까지 답안과 풀이 증거 제출 및 정산을 끝내야 합니다. 15:00까지 정산되지 않은 경기는 검토 대기로 전환되며 일요일 스냅샷을 바꾸지 않습니다.")
    callout(doc, "비경제적 보상", "배지·시즌 기록·Arena 활동 기록·프로필 장식은 학습일수로 바뀌지 않으며 페이백 자격을 만들지 않습니다.")

    return save(doc, "Matths_GOAT_Arena_공통_사용자_룰북_v1.0.docx")


def build_sub():
    doc = setup_document("GOAT Arena Sub Division 사용자 룰북", "Sub Division", "사용자용 v1.0")
    add_cover(
        doc,
        "Sub Division 규칙",
        "SUB DIVISION\n사용자 룰북",
        "학습권 · 일반 쟁탈전 · 복수전 · 페이백 · Main 진입",
        "사용자용 v1.0  |  Sub 활성 정책 v2.9",
        [
            "Sub Division의 목적과 핵심 용어",
            "결제주기·첫 학습일·일일 차감",
            "학습권 만료와 구매 제한",
            "일반 쟁탈전 신청 대상과 상대 선정",
            "일반 쟁탈전의 학습일수 정산",
            "Sub 복수전과 미완료 처리",
            "페이백 자격과 Main 진입",
            "Main 만료 뒤 Sub 재진입",
            "일요일 잠금과 시즌 운영",
        ],
    )

    chapter(doc, 1, "Sub Division 이해하기", "Sub Division은 Arena 랭킹, 학습일수를 이용한 1대1 경기, 결제주기별 페이백과 Main Division 진입 자격을 함께 운영합니다.")
    glossary(doc, [
        ("사용 가능 학습일수", "서비스 이용과 경기에 사용할 수 있는 현재 학습일수"),
        ("경기 예치 학습일수", "진행 중 경기의 정산이 끝날 때까지 사용할 수 없는 학습일수"),
        ("페이백 점수", "페이백 비율과 Main 진입 자격을 판단하는 결제주기별 점수"),
        ("Arena 상태", "Sub Arena 랭크, 랭크 안의 정확한 순위, Sub GP"),
        ("Skill MMR", "배치고사와 주간 공식 모의고사로만 바뀌는 시험 실력"),
        ("Final Ranking", "Skill MMR과 Division 성과 등을 함께 반영한 전체 종합 랭킹"),
    ])
    bullets(doc, [
        "Sub Division에는 무료 Arena 모드가 없습니다.",
        "공식 1대1 경기는 일반 쟁탈전과 복수전으로 구성됩니다.",
        "일반 쟁탈전과 복수전은 Skill MMR을 사용하거나 변경하지 않습니다.",
        "페이백과 Main 진입 자격은 결제주기별로 다시 판단합니다.",
    ])

    chapter(doc, 2, "결제주기와 첫 학습일", "정상 학습 패키지를 결제하면 새로운 결제주기가 시작되며 29일의 학습일수와 29점의 페이백 점수가 주어집니다.")
    heading(doc, "새 결제주기의 시작값")
    table(doc, ["항목", "시작 기준"], [
        ("사용 가능 학습일수", "29일"),
        ("페이백 점수", "29점"),
        ("완료한 유료 일반 공격", "0회"),
        ("연속 학습일", "0일"),
    ], [3300, 6060], font_size=9.4)
    bullets(doc, [
        "과거 결제주기의 남은 학습일수와 페이백 점수는 새 결제주기로 넘어오지 않습니다.",
        "최초 배치고사와 결제는 어느 순서로 완료해도 되지만, 둘 다 완료해야 활성 Sub Arena 순위를 사용할 수 있습니다.",
        "배치고사로 정해진 Skill MMR과 최초 Sub GP는 서로 다른 값으로 유지됩니다.",
    ])
    heading(doc, "결제 승인 시각에 따른 첫날 처리")
    table(doc, ["결제 승인 시각", "첫 학습일 처리"], [
        ("20:00 이전", "결제 당일을 첫 학습일로 계산하고 승인 시 학습일수 1일을 적용"),
        ("20:00부터 이후", "결제 당일은 차감하지 않고 다음 00:00부터 첫 학습일을 적용"),
    ], [2600, 6760], font_size=9.4)
    bullets(doc, [
        "정확히 20:00:00부터 이후 결제로 처리하며, 결제 화면이 아니라 서버의 승인 시각을 기준으로 합니다.",
        "20:00 이후 유예는 한 결제주기에 한 번만 적용되며, 일요일 잠금 중 결제한 Arena 이용은 월요일 00:00 이후 가능합니다.",
    ])

    chapter(doc, 3, "일일 차감·만료·추가 구매", "첫 학습일 이후에는 대한민국 표준시의 날짜가 바뀔 때마다 활성 결제주기의 사용 가능 학습일수가 1일씩 줄어듭니다.")
    heading(doc, "일일 차감")
    bullets(doc, [
        "일일 시간 차감은 사용 가능 학습일수에만 적용됩니다.",
        "일일 시간 차감만으로 페이백 점수가 줄어들지는 않습니다.",
        "진행 중 경기에 예치된 학습일수가 있으면 경기 정산을 먼저 진행합니다.",
    ])
    heading(doc, "새 패키지를 살 수 있는 시점")
    bullets(doc, [
        "사용 가능 학습일수가 0이어야 합니다.",
        "진행 중 경기에 예치된 학습일수가 없어야 합니다.",
        "아직 끝나지 않은 경기 정산이 없어야 합니다.",
        "학습일수가 남아 있는 동안에는 새 패키지를 미리 추가 구매할 수 없습니다.",
    ])
    heading(doc, "학습권 만료 뒤 기능")
    table(doc, ["구분", "기능"], [
        ("이용 제한", "Sub 공격·방어·복수전과 Main 경기, 주간 공식 모의고사·보너스, Skill MMR 갱신, 활성 Final Ranking, 새 페이백 점수"),
        ("계속 확인 가능", "결제 화면, 마지막 Arena·Final Rank, 과거 전적, 계정·결제 관리, 공지와 상세 규칙"),
    ], [2500, 6860], font_size=9.0)
    callout(doc, "방어 후보 제외", "학습일수가 만료된 사용자는 무료 방어자로 남지 않으며 신규 상대 후보에서 제외됩니다.")

    chapter(doc, 4, "일반 쟁탈전 신청", "Sub Division의 일반 쟁탈전에서는 사용자가 개인 상대가 아니라 목표 티어를 선택하고, 서버가 적격 상대 한 명을 무작위로 정합니다.")
    table(doc, ["도전자 티어", "신청 가능한 방어자 티어"], [
        ("브론즈", "브론즈 또는 실버"),
        ("실버", "골드"),
        ("골드", "플래티넘"),
        ("플래티넘", "에메랄드"),
        ("에메랄드", "다이아몬드"),
        ("다이아몬드", "마스터"),
        ("마스터", "그랜드마스터"),
        ("그랜드마스터", "챌린저"),
        ("챌린저", "챌린저"),
    ], [3600, 5760], font_size=9.4)
    bullets(doc, [
        "브론즈와 챌린저의 경계 규칙을 제외하면 자신의 바로 위 티어에만 신청할 수 있습니다.",
        "실버 이상 챌린저 미만 사용자는 같은 티어나 두 단계 이상 차이 나는 티어에 신청할 수 없습니다.",
        "선정된 방어자는 자동 참가하며 개인 상대를 직접 고를 수 없습니다.",
        "상대가 배정되기 전에는 개인 후보와 닉네임을 공개하지 않습니다.",
        "적격 상대가 없으면 경기와 학습일수 잠금 없이 신청이 종료됩니다.",
        "공식 매치가 성립한 뒤 양측은 24시간 안에 경기를 시작해야 하며, 일요일 14:30을 넘길 수 없습니다.",
        "경기에서는 자동 검산한 같은 주관식 준킬러 5문항을 10분 동안 풀며 자세한 제출·판정은 공통 룰북을 따릅니다.",
    ])

    chapter(doc, 5, "일반 쟁탈전 정산", "일반 쟁탈전을 신청하면 도전자의 학습일수 1일이 경기 정산이 끝날 때까지 잠깁니다. 방어자는 경기 생성 시 학습일수를 잠그지 않습니다.")
    table(doc, ["결과", "Arena 상태", "도전자의 1일", "방어자"], [
        ("실버 이상 도전자 승리", "양측 전체 교환", "소각", "학습일수 변화 없음"),
        ("브론즈 도전자 승리", "양측 전체 교환", "도전자에게 반환", "학습일수 변화 없음"),
        ("방어자 승리", "교환 없음", "방어자에게 이전", "학습일수·페이백 점수 각 1 증가"),
    ], [2350, 2050, 2350, 2610], font_size=8.7)
    bullets(doc, [
        "실버 이상 도전자가 승리하면 도전자의 사용 가능 학습일수와 페이백 점수가 각각 1 줄어듭니다.",
        "브론즈 도전자의 반환 규칙은 경기 시작 전 티어가 브론즈일 때만 적용합니다.",
        "방어자가 승리하면 도전자의 사용 가능 학습일수와 페이백 점수가 각각 1 줄고, 방어자의 두 값은 각각 1 늘어납니다.",
        "도전자가 승리하면 Arena 랭크·정확한 순위·GP를 양측이 전부 교환합니다.",
        "방어자가 승리하면 Arena 상태는 바뀌지 않습니다.",
    ])

    chapter(doc, 6, "Sub 복수전", "가장 최근 원경기의 패자에게 결과 화면에서 한 번의 복수전 선택권이 주어집니다.")
    bullets(doc, [
        "복수하기를 누르면 도전자의 학습일수 2일이 잠기고 상대는 자동 참가합니다.",
        "경기 종료를 선택하면 해당 원경기의 복수전 권리는 사라집니다.",
        "신청 뒤 24시간 안에 양측이 문제 풀이를 완료해야 하며 일요일 14:30을 넘길 수 없습니다.",
        "문제 형식, 승패 판정, 풀이 증거, Arena 상태 교환과 공정성 기준은 일반 쟁탈전과 같습니다.",
    ])
    table(doc, ["복수전 결과", "Arena 상태", "잠긴 2일 처리"], [
        ("도전자 정상 승리", "양측 전체 교환", "2일 전부 소각"),
        ("방어자 정상 승리", "교환 없음", "1일은 방어자에게 이전, 1일은 소각"),
        ("방어자만 미완료", "양측 전체 교환", "1일은 도전자에게 반환, 1일은 소각"),
        ("도전자만 미완료", "교환 없음", "1일은 방어자에게 이전, 1일은 소각"),
        ("양측 모두 미완료", "변경 보류", "자동 정산 없이 운영 검토 대기"),
    ], [2900, 2350, 4110], font_size=8.9)
    callout(doc, "양측 미완료", "두 사용자 모두 기한 안에 완료하지 않으면 Arena 상태와 학습일수를 자동으로 바꾸지 않고 운영 검토 대기로 전환합니다.")

    chapter(doc, 7, "페이백과 Main 진입", "페이백은 각 결제주기의 연속 학습, 유료 일반 공격, 페이백 점수와 공정성 상태를 함께 확인해 결정합니다.")
    heading(doc, "페이백 자격")
    bullets(doc, [
        "연속 학습일이 30일 이상이어야 합니다.",
        "유료 일반 공격을 2회 이상 완료해야 합니다.",
        "페이백 점수가 30점 이상이어야 합니다.",
        "부정행위 또는 공모 관련 문제가 없는 정상 상태여야 합니다.",
    ])
    heading(doc, "페이백 구간")
    table(doc, ["최종 페이백 점수", "페이백 비율"], [
        ("29점 이하", "0%"),
        ("30~34점", "50%"),
        ("35~39점", "80%"),
        ("40점 이상", "100%"),
    ], [5000, 4360], font_size=9.6)
    bullets(doc, [
        "각 결제주기는 정해진 평가 시점에 한 번만 심사합니다.",
        "사용자마다 결제 시점이 다르므로 심사 시점도 각 결제주기를 기준으로 합니다.",
        "새 결제주기가 시작되어도 이전 결제주기의 심사는 별도로 진행됩니다.",
        "페이백 자격이 확정되면 실제 송금 완료 전에도 Main에 진입할 수 있습니다.",
    ])
    heading(doc, "Main 시작 학습일수")
    para(doc, "Main에는 최종 Sub 페이백 점수 중 기본 29점을 넘긴 부분을 학습일수로 이월하고, 같은 Sub 결제주기에서 한 번만 Main 진입 보너스 2일을 더합니다. Sub의 페이백 점수 전체를 Main 잔액으로 옮기지는 않으며, Main 진입 뒤에는 페이백 점수를 새로 쌓거나 Main 페이백 심사를 다시 진행하지 않습니다.")

    chapter(doc, 8, "Main 만료 뒤 Sub 재진입", "Main의 모든 학습일수가 끝나고 예약·잠금·미정산 경기가 없으면 사용자는 Sub로 강등되며, 재구독 전까지 Arena와 주간 공식 모의고사를 이용할 수 없습니다.")
    heading(doc, "강등 시 보존되는 기록")
    bullets(doc, [
        "마지막 Main 랭크·정확한 순위·GP",
        "마지막 Main 백분위와 참가자 수",
        "프로필의 Main 달성 이력",
        "과거 경기와 학습일수 기록",
    ])
    heading(doc, "72시간 안에 재구독")
    bullets(doc, [
        "별도 시험 없이 직전 Main 성과를 반영한 Sub 랭크에서 시작합니다.",
        "새 Sub 결제주기는 29일의 학습일수와 29점의 페이백 점수로 시작합니다.",
        "현재 경쟁 Division은 Sub이며 과거 Main 달성 기록은 유지됩니다.",
        "새 결제주기에서 페이백 조건을 다시 충족하면 Main에 다시 진입할 수 있습니다.",
    ])
    heading(doc, "72시간이 지난 뒤 재구독")
    bullets(doc, [
        "랭크 복귀전을 완료해야 Sub Arena가 활성화됩니다.",
        "랭크 복귀전은 늦은 재구독자의 Sub 재진입 랭크만 결정하며 Skill MMR을 초기화하지 않습니다.",
        "시험 성적이 높아도 72시간 안에 재구독했을 때 받을 기준 위치보다 낮은 Sub 위치에서 시작합니다.",
        "시험을 마치기 전에는 일반 학습은 가능하지만 Arena·주간 공식 모의고사·활성 Final Ranking은 이용할 수 없습니다.",
    ])
    heading(doc, "Sub 사용자가 만료된 뒤")
    bullets(doc, [
        "학습일수가 0이면 Arena와 주간 공식 모의고사를 이용할 수 없습니다.",
        "새 패키지는 기존 학습일수가 0일 때만 살 수 있으며, 새 결제주기에는 새로운 페이백 기회가 주어지지만 이전 결제주기의 페이백 점수는 이월되지 않습니다.",
    ])

    chapter(doc, 9, "일요일 잠금", "매주 일요일에는 주간 공식 모의고사와 랭킹 갱신을 위해 Sub와 Main Arena를 함께 잠급니다.")
    table(doc, ["시간", "Sub Division 처리"], [
        ("일요일 14:30부터", "신규 경기 신청·수락·준비·시작 차단"),
        ("일요일 15:00부터", "Arena 경기 기록·정산 잠금 및 공개 Final Ranking 고정"),
        ("15:00~24:00", "정상 이용자의 공식 모의고사와 새 랭킹 반영 준비"),
        ("월요일 00:00", "새 MMR·주간 보너스·Final Rank 공개 및 Arena 잠금 해제"),
    ], [2600, 6760], font_size=9.2)
    bullets(doc, [
        "14:30 전에 시작한 경기는 15:00까지 답안·풀이 증거 제출과 정산을 끝내야 합니다.",
        "15:00까지 끝나지 않은 경기는 운영 검토 대기로 전환합니다.",
        "학습권이 만료된 사용자는 주간 공식 모의고사와 새 랭킹 반영 대상에서 제외됩니다.",
    ])

    return save(doc, "Matths_GOAT_Arena_Sub_Division_사용자_룰북_v1.0.docx")


def build_main():
    doc = setup_document("GOAT Arena Main Division 사용자 룰북", "Main Division", "사용자용 v1.0")
    add_cover(
        doc,
        "Main Division 규칙",
        "MAIN DIVISION\n사용자 룰북",
        "진입 · 학습일수 · 일반 상향 공격 · 초대전 · 복수전",
        "사용자용 v1.0  |  Main 운영 정책 v1.1",
        [
            "Main 진입과 시작 학습일수",
            "학습일수의 사용 가능·예약·경기 예치 상태",
            "예치 범위와 상대 선정",
            "일반 상향 공격",
            "상위 티어의 하위 티어 초대전",
            "초대 예약·취소·자동 취소",
            "Main 복수전",
            "경기 횟수·중복 제한·일요일 잠금",
            "학습일수 만료·Sub 강등·재구독",
        ],
    )

    chapter(doc, 1, "Main Division 이해하기", "Main Division은 Sub Division에서 페이백과 Main 진입 조건을 달성한 사용자가 남은 학습일수를 예치하고 경쟁하는 상위 Arena입니다.")
    glossary(doc, [
        ("사용 가능 학습일수", "현재 자유롭게 이용하거나 경기에 예치할 수 있는 학습일수"),
        ("예약 중 학습일수", "상위→하위 초대가 수락되기 전까지 따로 보관된 학습일수"),
        ("경기 예치 학습일수", "매치가 성립한 뒤 경기 정산까지 사용할 수 없는 학습일수"),
        ("Arena 상태", "Main Arena 랭크, 랭크 안의 정확한 순위, Main GP"),
        ("일반 상향 공격", "하위 티어 사용자가 더 높은 티어에 도전하는 경기"),
        ("상위→하위 초대전", "상위 티어 사용자가 하위 티어 사용자에게 선택 가능한 경기 기회를 보내는 방식"),
    ])
    bullets(doc, [
        "Main에서는 페이백 점수를 새로 쌓거나 페이백 심사를 다시 진행하지 않습니다.",
        "Main 경기 결과는 Skill MMR을 바꾸지 않습니다.",
        "Skill MMR은 배치고사와 매주 일요일 공식 모의고사에서만 바뀝니다.",
        "Sub와 Main의 참가자·Arena 상태·상대 선정·예치 규칙은 서로 분리됩니다.",
    ])

    chapter(doc, 2, "Main 진입과 학습일수", "Sub에서 페이백과 Main 진입 자격을 모두 달성하면 Main 소속과 시작 학습일수를 받습니다.")
    heading(doc, "Main 시작 학습일수")
    bullets(doc, [
        "최종 Sub 페이백 점수 중 기본 29점을 넘긴 부분을 Main 학습일수로 이월합니다.",
        "같은 Sub 결제주기에서 한 번만 Main 진입 보너스 2일을 더합니다.",
        "중복 결제 알림이나 중복 심사가 발생해도 같은 보너스를 다시 지급하지 않습니다.",
    ])
    heading(doc, "학습일수의 출처")
    table(doc, ["출처", "설명"], [
        ("Sub 이월분", "최종 Sub 페이백 점수에서 기본 29점을 제외한 학습일수"),
        ("Main 진입 보너스", "Main에 처음 진입할 때 같은 Sub 결제주기에서 한 번 받는 2일"),
        ("Main 경기 이전분", "Main 경기 승리로 다른 사용자에게서 이전받은 학습일수"),
    ], [2800, 6560], font_size=9.3)
    heading(doc, "일일 차감 순서")
    numbered(doc, [
        "Sub에서 이월한 학습일수",
        "Main 진입 보너스 학습일수",
        "Main 경기에서 이전받은 학습일수",
    ])
    para(doc, "대한민국 표준시의 날짜가 바뀔 때마다 활성 Main 사용자의 학습일수 1일이 차감됩니다. 학습일수가 남아 있으면 매주 공식 모의고사에 응시할 수 있습니다.")

    chapter(doc, 3, "사용 가능·예약·경기 예치 상태", "Main의 전체 학습일수는 사용 가능 학습일수, 예약 중 학습일수와 경기 예치 학습일수로 나뉩니다.")
    table(doc, ["상태", "사용자가 할 수 있는 일"], [
        ("사용 가능", "서비스 이용과 새 경기 예치에 사용 가능"),
        ("예약 중", "초대 수락 전까지 다른 경기에서 중복 사용 불가"),
        ("경기 예치", "공식 매치 정산이 끝날 때까지 이동·취소·재사용 불가"),
    ], [2500, 6860], font_size=9.4)
    bullets(doc, [
        "화면의 총 Main 학습일수에는 세 상태의 학습일수가 모두 포함됩니다.",
        "새 경기에 사용할 수 있는 잔액은 사용 가능 학습일수만을 기준으로 합니다.",
        "예약과 예치는 그 자체로 학습일수 소비가 아닙니다.",
        "정상 취소나 경기 무효 처리 시 Division 규칙에 따라 사용 가능 학습일수로 반환합니다.",
    ])
    callout(doc, "중복 사용 금지", "예약 중이거나 경기에 예치된 학습일수는 다른 공격·초대전·복수전에 다시 예치할 수 없습니다.")

    chapter(doc, 4, "예치와 상대 선정", "Main에서는 개인 상대를 직접 고르지 않고 목표 티어와 예치할 학습일수를 선택합니다. 서버가 적격 후보 중 한 명을 무작위로 정합니다.")
    heading(doc, "티어 차이별 최소 예치 일수")
    table(doc, ["목표 티어 차이", "최소 예치 일수", "신청 가능 여부"], [
        ("1단계", "1일", "가능"),
        ("2단계", "2일", "가능"),
        ("3단계", "3일", "가능"),
        ("4단계 이상", "-", "신청 불가"),
    ], [3300, 2700, 3360], font_size=9.5)
    bullets(doc, [
        "사용자는 최소 예치 일수 이상인 정수 학습일수를 선택할 수 있습니다.",
        "신청자와 상대 모두 예치를 마친 뒤 사용할 학습일수가 최소 1일 남아야 합니다.",
        "예치 일수와 사용 가능 학습일수가 같은 사용자는 신규 경기의 신청자나 상대 후보가 될 수 없습니다.",
        "부족한 학습일수를 운영자가 새로 보충해 매치를 성립시키지 않습니다.",
    ])
    heading(doc, "상대 후보에서 제외되는 경우")
    bullets(doc, [
        "요청자 본인 또는 공식적으로 연관된 계정",
        "정지·제재·공정성 검토 중인 계정",
        "다른 공식 경기의 정산이 끝나지 않은 사용자",
        "반복 매칭 제한 대상",
        "예치 뒤 최소 1일을 남길 수 없는 사용자",
        "현재 시즌 배치 절차를 완료하지 않았거나 일요일 잠금 중인 사용자",
    ])

    chapter(doc, 5, "일반 상향 공격", "낮은 티어 사용자가 높은 티어를 선택하면 서버가 해당 티어의 적격 방어자 한 명을 무작위로 선정합니다. 선정된 방어자는 공식 경기 규칙에 따라 자동 참가합니다.")
    bullets(doc, [
        "공격자와 방어자는 같은 학습일수를 각각 경기 종료까지 예치합니다.",
        "두 사용자 모두 예치 뒤 사용할 학습일수가 최소 1일 남아야 매치가 성립합니다.",
        "문제·채점·풀이 증거·승패 판정은 공통 룰북을 따릅니다.",
    ])
    table(doc, ["결과", "Arena 상태", "학습일수 정산"], [
        ("공격자 승리", "양측 전체 교환", "공격자는 자신의 예치분을 돌려받고 방어자의 예치분을 받음"),
        ("방어자 승리", "교환 없음", "방어자는 자신의 예치분을 돌려받고 공격자의 예치분을 받음"),
    ], [2300, 2350, 4710], font_size=9.1)
    para(doc, "도전자가 승리할 때만 두 사용자의 Arena 랭크·정확한 순위·GP가 전부 교환됩니다. 방어자가 승리하면 Arena 상태는 그대로 유지됩니다.")

    chapter(doc, 6, "상위 티어의 하위 티어 초대전", "상위 티어 사용자는 자신의 Arena 상태를 걸고 학습일수를 예치해 하위 티어 사용자에게 경기 기회를 보낼 수 있습니다.")
    heading(doc, "초대 방식")
    bullets(doc, [
        "상위 사용자는 개인 상대가 아니라 목표 하위 티어를 선택합니다.",
        "서버가 해당 티어의 적격 후보를 무작위로 선정해 초대를 보냅니다.",
        "하위 사용자는 초대를 수락하거나 거절할 수 있습니다.",
        "거절해도 Arena 상태, 학습일수와 Final Ranking에 불이익이 없습니다.",
        "가장 먼저 수락 절차를 완료한 한 명과만 매치를 만들고 나머지 초대는 종료합니다.",
    ])
    heading(doc, "경기에서의 역할")
    table(doc, ["사용자", "초대에서의 역할", "Arena 판정 역할"], [
        ("상위 티어 사용자", "초대 생성자", "방어자"),
        ("하위 티어 사용자", "초대 수신자", "수락 후 도전자"),
    ], [2700, 3300, 3360], font_size=9.4)
    heading(doc, "수락할 때")
    bullets(doc, [
        "서버가 양측의 계정·Division·티어·학습일수·일요일 잠금·미정산 경기 여부를 다시 확인합니다.",
        "상위 사용자의 예약 학습일수는 경기 예치 상태로 바뀝니다.",
        "하위 사용자는 같은 학습일수를 사용 가능 잔액에서 경기 예치 상태로 옮깁니다.",
        "조건을 충족하지 못하면 공식 매치를 만들지 않습니다.",
    ])
    heading(doc, "초대전 정산")
    para(doc, "상위 사용자가 이기면 Arena 상태는 바뀌지 않고 상위 사용자가 양측의 예치 학습일수를 받습니다. 하위 사용자가 이기면 양측의 Arena 상태가 전체 교환되고 하위 사용자가 양측의 예치 학습일수를 받습니다.")

    chapter(doc, 7, "초대 예약과 취소", "상위→하위 초대는 상대가 수락하기 전까지 공식 매치가 아니라 예약 상태로 유지됩니다.")
    heading(doc, "예약이 유지되는 기간")
    bullets(doc, [
        "초대 예약에는 고정 24시간 만료가 없습니다.",
        "초대 생성자가 취소하거나, 적격 사용자가 수락하거나, 생성자가 경기 자격을 잃거나, 자동 취소 조건이 발생하거나, 운영 검토로 무효가 될 때까지 유지됩니다.",
        "한 후보가 거절하거나 자격을 잃어도 예약은 유지되며 서버가 같은 목표 티어의 새로운 후보를 찾습니다.",
        "초대 생성자와 최근 7일 안에 공식 매치가 성립했던 사용자는 후보에서 제외됩니다.",
        "같은 초대 생성자는 같은 목표 티어에 미성립 예약을 동시에 하나만 유지할 수 있습니다.",
    ])
    heading(doc, "매치 성립 전 직접 취소")
    bullets(doc, [
        "초대 생성자는 매치가 성립하기 전 예약을 직접 취소할 수 있습니다.",
        "직접 취소 수수료는 없으며 예약 학습일수를 전부 사용 가능 학습일수로 돌려받습니다.",
        "매치가 성립해 학습일수가 경기 예치 상태가 된 뒤에는 임의로 취소할 수 없습니다.",
    ])
    heading(doc, "일일 차감으로 인한 자동 취소")
    bullets(doc, [
        "일일 차감은 먼저 사용 가능 학습일수에서 처리합니다.",
        "사용 가능 학습일수가 0이 되었는데 아직 성립되지 않은 초대 예약이 남아 있으면 예약을 자동 취소합니다.",
        "자동 취소 시 예약 학습일수 중 1일은 수수료로 소각하고 나머지를 사용 가능 학습일수로 반환합니다.",
        "예약 학습일수가 1일이면 반환되는 학습일수 없이 1일 전부가 수수료로 소각됩니다.",
        "이미 매치가 성립해 경기 예치 상태가 된 학습일수는 자동 취소하지 않고 경기 정산을 먼저 진행합니다.",
    ])

    chapter(doc, 8, "Main 복수전", "가장 최근 원경기의 패자만 결과 화면에서 복수하기를 선택할 수 있으며, 상대는 거절할 수 없습니다.")
    bullets(doc, [
        "경기 종료를 선택하면 해당 원경기의 복수전 권리는 즉시 사라집니다.",
        "복수전 신청자는 직전 원경기에서 양측이 걸었던 학습일수의 두 배를 잠급니다.",
        "신청 뒤 24시간 안에 양측이 문제 풀이를 완료해야 하며 일요일 14:30을 넘길 수 없습니다.",
        "문제 형식, 승패 우선순위, 완전 동점 시 방어자 승리, 증거 제출과 공정성 기준은 공통 룰북을 따릅니다.",
    ])
    table(doc, ["복수전 결과", "Arena 상태", "신청자가 잠근 학습일수 처리"], [
        ("도전자 정상 승리", "양측 전체 교환", "전부 소각"),
        ("방어자 정상 승리", "교환 없음", "1일 소각, 나머지는 방어자에게 이전"),
        ("방어자만 미완료", "양측 전체 교환", "1일 소각, 나머지는 도전자에게 반환"),
        ("도전자만 미완료", "교환 없음", "1일 소각, 나머지는 방어자에게 이전"),
        ("양측 모두 미완료", "교환 없음", "전부 소각"),
    ], [2850, 2250, 4260], font_size=8.8)

    chapter(doc, 9, "횟수와 중복 제한", "Main에는 일일 최대 공격·방어·초대전 횟수와 결제주기당 경기 순증가 학습일수 상한이 없습니다.")
    bullets(doc, [
        "한 사용자는 같은 시점에 정산되지 않은 공식 경기를 하나만 가질 수 있습니다.",
        "일일 횟수 제한이 없다는 것이 같은 학습일수를 여러 경기에 중복 사용할 수 있다는 뜻은 아닙니다.",
        "같은 목표 티어에 유지할 수 있는 미성립 초대 예약은 동시에 하나입니다.",
        "공식 매치가 성립했던 같은 상대는 7일 동안 다시 후보가 될 수 없습니다.",
        "예약 총액과 사용 가능 학습일수는 매치가 성립하기 전에 서버가 다시 확인합니다.",
    ])
    heading(doc, "공정성 검토")
    bullets(doc, [
        "연관 기기·네트워크·결제수단을 사용하는 계정끼리의 반복 경기를 확인합니다.",
        "한 방향으로 반복되는 학습일수 이전과 의도적인 오답·비정상 제출을 확인합니다.",
        "반복적인 초대 수락·거절을 통한 상대 탐색과 여러 계정을 이용한 학습일수 몰아주기를 확인합니다.",
        "의심 경기는 운영 검토 대기로 전환하며, 검토가 끝날 때까지 Arena 상태와 학습일수를 확정하지 않습니다.",
    ])

    chapter(doc, 10, "일요일 잠금", "Main의 기존 미성립 초대 예약은 일요일 잠금 중에도 취소되지 않지만, 새 후보 선정과 수락·매치 성립은 멈춥니다.")
    table(doc, ["시간", "Main Division 처리"], [
        ("일요일 14:30부터", "신규 공식 경기 매칭·초대 수락·준비·시작 차단"),
        ("일요일 15:00부터", "Arena 상태를 바꾸는 경기 기록·정산 잠금 및 공개 Final Ranking 고정"),
        ("14:30~월요일 00:00", "미성립 초대 예약 유지, 신규 후보 선정·수락·매치 성립 중단"),
        ("월요일 00:00", "자격·잔액·티어를 다시 확인한 뒤 적격 초대 매칭 재개"),
    ], [2700, 6660], font_size=9.1)
    bullets(doc, [
        "정상 Main 사용자는 잠금 시간 동안 주간 공식 모의고사에 응시할 수 있습니다.",
        "새 Skill MMR·주간 보너스·Final Rank는 월요일 00:00에 함께 공개됩니다.",
    ])

    chapter(doc, 11, "만료·Sub 강등·재구독", "사용 가능·예약 중·경기 예치 학습일수가 모두 0이고 미정산 경기가 없으면 Main 이용이 끝나며 사용자는 Sub로 강등됩니다.")
    heading(doc, "만료 전에 먼저 처리되는 항목")
    bullets(doc, [
        "사용 가능 학습일수가 0이어도 미성립 초대 예약이 남아 있으면 예약 자동 취소 규칙을 먼저 적용합니다.",
        "경기에 예치된 학습일수가 있으면 해당 경기 정산을 먼저 진행합니다.",
        "모든 학습일수와 미정산 경기가 정리된 뒤 최종 만료를 확정합니다.",
    ])
    heading(doc, "강등 뒤 보존되는 기록")
    bullets(doc, [
        "마지막 Main 랭크·정확한 순위·GP와 백분위",
        "당시 Main 참가자 수",
        "Main 성취 이력과 프로필 배지",
        "Main 경기와 학습일수 기록",
    ])
    heading(doc, "강등 뒤 제한")
    bullets(doc, [
        "현재 경쟁 Division은 Sub로 변경됩니다.",
        "재구독 전까지 GOAT Arena와 주간 공식 모의고사를 이용할 수 없습니다.",
        "활성 Main 상대 후보·초대 생성·활성 Final Ranking 자격이 종료됩니다.",
    ])
    heading(doc, "재구독")
    table(doc, ["재구독 시점", "Sub 재진입 방식"], [
        ("만료 후 72시간 안", "시험 없이 직전 Main 성과를 반영한 Sub 랭크에서 새 결제주기 시작"),
        ("만료 후 72시간 초과", "랭크 복귀전을 완료하고 정상 갱신 기준보다 낮은 Sub 위치에서 시작"),
    ], [3200, 6160], font_size=9.2)
    bullets(doc, [
        "재구독하더라도 Main에서 바로 이어서 시작하지 않습니다.",
        "새 Sub 결제주기는 29일의 학습일수와 29점의 페이백 점수로 시작합니다.",
        "새 Sub 결제주기에서 페이백과 Main 진입 조건을 다시 달성하면 Main에 재진입할 수 있습니다.",
        "72시간이 지난 뒤 치르는 랭크 복귀전은 Skill MMR을 초기화하지 않습니다.",
    ])
    heading(doc, "연간 시즌")
    bullets(doc, [
        "연도가 바뀌어도 Main 달성 이력과 소속 이력은 유지합니다.",
        "새 시즌에는 Arena 랭크·정확한 순위·GP와 Final Rank를 초기화합니다.",
        "Main 내부 시즌 배치고사를 완료해야 새 시즌 Arena와 Final Ranking에 참여할 수 있습니다.",
        "시즌 배치고사와 늦은 재구독자의 랭크 복귀전은 서로 다른 시험입니다.",
    ])

    return save(doc, "Matths_GOAT_Arena_Main_Division_사용자_룰북_v1.0.docx")


def main():
    paths = [build_common(), build_sub(), build_main()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
